
# relatorio_fotografico/services/docx_generator.py
"""
Gerador de relatório fotográfico em Word (.docx).

As imagens usadas aqui já estão sanitizadas e padronizadas (tamanho
fixo, JPEG otimizado) desde o upload — ver `FotoRelatorio.save()` em
models.py. Nenhum processamento de imagem é feito neste módulo.
"""
import io
import os

from django.conf import settings

from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn

from ..models import FOTOS_POR_PAGINA

# --- Layout (cabeçalho) ---
LARGURA_LOGO_COL = Cm(3)
LARGURA_VAZIO_COL = Cm(3)
LARGURA_LOGO_IMG = Cm(2.5)

# --- Layout (grid de fotos) ---
LINHAS_GRID = 3
COLUNAS_GRID = 2
LARGURA_IMAGEM = Cm(7.5)
ALTURA_IMAGEM = Cm(5.6)
ALTURA_LINHA = Cm(6.5)  # imagem + espaço da legenda

MARGEM_ESQUERDA = Cm(1.5)
MARGEM_DIREITA = Cm(1.5)
MARGEM_SUPERIOR = Cm(1.2)
MARGEM_INFERIOR = Cm(1.2)


def _resolver_logo_path():
    """
    Resolve o caminho físico de static/images/logocetest.png,
    funcionando tanto em dev (STATICFILES_DIRS) quanto após
    collectstatic (STATIC_ROOT).
    """
    candidatos = []

    if getattr(settings, 'STATIC_ROOT', None):
        candidatos.append(
            os.path.join(settings.STATIC_ROOT, 'images', 'logocetest.png')
        )

    for static_dir in getattr(settings, 'STATICFILES_DIRS', []):
        candidatos.append(os.path.join(static_dir, 'images', 'logocetest.png'))

    for caminho in candidatos:
        if os.path.exists(caminho):
            return caminho
    return None


def _set_cell_border(cell, color="000000", sz=8):
    """Aplica bordas manuais em uma célula (não usado no fluxo atual,
    mantido como utilitário caso a tabela precise de bordas customizadas
    sem depender do estilo 'Table Grid')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.makeelement(qn('w:tcBorders'), {})
    for edge in ('top', 'left', 'bottom', 'right'):
        el = tc_borders.makeelement(qn(f'w:{edge}'), {
            qn('w:val'): 'single',
            qn('w:sz'): str(sz),
            qn('w:color'): color,
        })
        tc_borders.append(el)
    tc_pr.append(tc_borders)


def _resolver_caminho_foto(foto):
    """Caminho físico da imagem (já padronizada no upload)."""
    return foto.imagem.path


def _configurar_secao(doc):
    """Define margens e retorna a largura útil da página."""
    section = doc.sections[0]
    section.left_margin = MARGEM_ESQUERDA
    section.right_margin = MARGEM_DIREITA
    section.top_margin = MARGEM_SUPERIOR
    section.bottom_margin = MARGEM_INFERIOR
    return section.page_width - section.left_margin - section.right_margin


def _adicionar_cabecalho_titulo(doc, largura_util):
    """Tabela 1x3: logo | título centralizado | espaço vazio (simetria)."""
    top_table = doc.add_table(rows=1, cols=3)
    top_table.autofit = False

    largura_titulo = largura_util - LARGURA_LOGO_COL - LARGURA_VAZIO_COL

    top_table.columns[0].width = LARGURA_LOGO_COL
    top_table.columns[1].width = largura_titulo
    top_table.columns[2].width = LARGURA_VAZIO_COL

    cel_logo, cel_titulo, cel_vazio = top_table.rows[0].cells
    cel_logo.width = LARGURA_LOGO_COL
    cel_titulo.width = largura_titulo
    cel_vazio.width = LARGURA_VAZIO_COL

    logo_path = _resolver_logo_path()
    if logo_path:
        run_logo = cel_logo.paragraphs[0].add_run()
        run_logo.add_picture(logo_path, width=LARGURA_LOGO_IMG)

    p_titulo = cel_titulo.paragraphs[0]
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_titulo = p_titulo.add_run('RELATÓRIO FOTOGRÁFICO')
    run_titulo.bold = True
    run_titulo.font.size = Pt(16)


def _adicionar_cabecalho_dados(doc, relatorio):
    """Tabela 3x2 com dados do relatório (obra, data, assunto, folha, responsável)."""
    header = doc.add_table(rows=3, cols=2)
    header.style = 'Table Grid'
    header.alignment = WD_TABLE_ALIGNMENT.CENTER

    header.cell(0, 0).text = f"Obra/Contrato: {relatorio.obra_contrato}"
    header.cell(0, 1).text = f"Data: {relatorio.data:%d/%m/%Y}"
    header.cell(1, 0).text = f"Assunto: {relatorio.titulo}"
    header.cell(1, 1).text = f"Folha: 01 de {relatorio.total_folhas:02d}"

    responsavel_nome = (
        relatorio.responsavel.get_full_name()
        or relatorio.responsavel.username
    )
    header.cell(2, 0).merge(header.cell(2, 1))
    header.cell(2, 0).text = f"Responsável: {responsavel_nome}"


def _adicionar_titulo_pagina(doc, relatorio, pagina_idx, total_folhas):
    """Cabeçalho de continuação, exibido a partir da 2ª página."""
    doc.add_page_break()
    p = doc.add_paragraph()
    run = p.add_run(
        f"{relatorio.obra_contrato} — {relatorio.titulo} "
        f"— Folha {pagina_idx:02d} de {total_folhas:02d}"
    )
    run.bold = True
    run.font.size = Pt(10)


def _preencher_celula_foto(cell, foto, numero_foto):
    """Insere imagem + legenda numerada em uma célula do grid."""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    try:
        caminho_imagem = _resolver_caminho_foto(foto)
        # Largura E altura fixas — imagens já vêm padronizadas em 4:3
        # desde o upload (ImageOps.fit em FotoRelatorio._padronizar_imagem).
        run.add_picture(
            caminho_imagem,
            width=LARGURA_IMAGEM,
            height=ALTURA_IMAGEM,
        )
    except Exception:
        paragraph.add_run('[imagem indisponível]')
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    legenda_p = cell.add_paragraph()
    legenda_run = legenda_p.add_run(f"Foto {numero_foto:02d}: {foto.legenda}")
    legenda_run.font.size = Pt(9)
    legenda_p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def _adicionar_grid_fotos(doc, fotos_pagina, pagina_idx, largura_util):
    """Monta a tabela 3x2 (grid) de fotos de uma página."""
    largura_coluna = int(largura_util / COLUNAS_GRID)

    table = doc.add_table(rows=LINHAS_GRID, cols=COLUNAS_GRID)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    # Fixa a largura de todas as colunas/células (evita caixas desiguais)
    for col in table.columns:
        col.width = largura_coluna
    for row in table.rows:
        for cell in row.cells:
            cell.width = largura_coluna

    for idx in range(LINHAS_GRID * COLUNAS_GRID):
        row = idx // COLUNAS_GRID
        col = idx % COLUNAS_GRID
        cell = table.cell(row, col)

        if idx < len(fotos_pagina):
            foto = fotos_pagina[idx]
            numero_foto = idx + 1 + (pagina_idx - 1) * FOTOS_POR_PAGINA
            _preencher_celula_foto(cell, foto, numero_foto)
        else:
            cell.text = ''

    # Trava a altura de todas as linhas desta tabela
    for row in table.rows:
        row.height = ALTURA_LINHA
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY


def _adicionar_rodape(doc):
    """Rodapé fixo, aplicado em todas as seções do documento."""
    for sec in doc.sections:
        footer_p = sec.footer.paragraphs[0]
        footer_p.text = ''
        footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_footer = footer_p.add_run('Relatório Fotográfico — CETEST')
        run_footer.font.size = Pt(9)
        run_footer.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def gerar_docx_relatorio(relatorio):
    """
    Gera o relatório em Word (.docx).

    As imagens de cada `FotoRelatorio` já estão sanitizadas e
    padronizadas (tamanho fixo, JPEG otimizado) desde o upload —
    são usadas diretamente aqui, sem processamento adicional.

    Args:
        relatorio: instância de RelatorioFotografico.

    Returns:
        io.BytesIO: buffer com o conteúdo do .docx gerado.
    """
    doc = Document()

    largura_util = _configurar_secao(doc)

    _adicionar_cabecalho_titulo(doc, largura_util)
    _adicionar_cabecalho_dados(doc, relatorio)
    doc.add_paragraph()

    paginas = relatorio.paginas
    total_folhas = relatorio.total_folhas

    for pagina_idx, fotos_pagina in enumerate(paginas, start=1):
        if pagina_idx > 1:
            _adicionar_titulo_pagina(doc, relatorio, pagina_idx, total_folhas)

        _adicionar_grid_fotos(doc, fotos_pagina, pagina_idx, largura_util)

    _adicionar_rodape(doc)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer



