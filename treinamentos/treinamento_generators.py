
import io
from datetime import datetime
import pandas as pd # Mantido para contexto, mas não usado na função Word
from django.http import HttpResponse # Mantido para contexto
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# --- Imports para Gráficos ---
import matplotlib
matplotlib.use('Agg') # Usa um backend não-interativo, essencial para apps web
import matplotlib.pyplot as plt


def gerar_relatorio_word(treinamento, caminho_logo=None):
    """
    Gera um relatório .docx robusto e visual para um treinamento,
    incluindo logomarca, análises detalhadas e gráficos.

    :param treinamento: O objeto do modelo Treinamento.
    :param caminho_logo: (Opcional) O caminho para o arquivo de imagem da logomarca.
    :return: Um buffer de memória com o arquivo .docx pronto para a HttpResponse.
    """
    document = Document()

    # ==================================================================
    # 1. PRÉ-PROCESSAMENTO DOS DADOS
    # ==================================================================
    # Coleta todos os dados necessários de uma vez para evitar múltiplas queries.
    participantes_queryset = treinamento.participantes.select_related('funcionario').all()
    lista_participantes = list(participantes_queryset)

    total_inscritos = len(lista_participantes)
    participantes_presentes = [p for p in lista_participantes if p.presente]
    total_presentes = len(participantes_presentes)
    total_ausentes = total_inscritos - total_presentes
    taxa_presenca = (total_presentes / total_inscritos * 100) if total_inscritos > 0 else 0

    # Análise de Custos
    custo_por_participante_presente = (treinamento.custo / total_presentes) if total_presentes > 0 else 0

    # Análise de Desempenho (considerando apenas os que têm nota)
    notas = [p.nota_avaliacao for p in participantes_presentes if p.nota_avaliacao is not None]
    media_notas = sum(notas) / len(notas) if notas else 0
    maior_nota = max(notas) if notas else 0
    menor_nota = min(notas) if notas else 0
    # Supondo que a nota para aprovação seja 7.0
    aprovados = [n for n in notas if n >= 7.0]
    taxa_aprovacao = (len(aprovados) / len(notas) * 100) if notas else 0


    # ==================================================================
    # 2. CONFIGURAÇÃO DO DOCUMENTO (CABEÇALHO E ESTILOS)
    # ==================================================================
    # --- Cabeçalho com Logomarca ---
    if caminho_logo:
        section = document.sections[0]
        header = section.header
        p_header = header.paragraphs[0]
        p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_header = p_header.add_run()
        # Adiciona a imagem com uma altura fixa, a largura será ajustada proporcionalmente
        run_header.add_picture(caminho_logo, height=Cm(1.5))

    # --- Título Principal ---
    p_title = document.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = p_title.add_run('RELATÓRIO DE TREINAMENTO')
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(18)
    run_title.font.bold = True
    document.add_paragraph() # Espaçamento

    # ==================================================================
    # 3. CONTEÚDO DO RELATÓRIO
    # ==================================================================

    # --- Seção 1: Informações Gerais ---
    document.add_heading('1. Informações Gerais do Treinamento', level=1)
    table_info = document.add_table(rows=0, cols=2)
    table_info.style = 'Table Grid'
    table_info.columns[0].width = Inches(2.5)
    table_info.columns[1].width = Inches(4.0)

    # Helper para adicionar linhas à tabela
    def add_info_row(table, label, value):
        cells = table.add_row().cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].font.bold = True
        cells[1].text = str(value)

    add_info_row(table_info, 'Nome do Treinamento:', treinamento.nome)
    add_info_row(table_info, 'Tipo de Curso:', treinamento.tipo_curso.nome)
    add_info_row(table_info, 'Status:', treinamento.get_status_display())
    add_info_row(table_info, 'Palestrante / Instrutor:', treinamento.palestrante)
    add_info_row(table_info, 'Local:', treinamento.local)
    add_info_row(table_info, 'Período:', f"{treinamento.data_inicio.strftime('%d/%m/%Y')} a {treinamento.data_vencimento.strftime('%d/%m/%Y')}")
    add_info_row(table_info, 'Carga Horária:', f"{treinamento.duracao} horas")
    add_info_row(table_info, 'Custo Total:', f"R$ {treinamento.custo:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'))


    # --- Seção 2: Análise de Resultados ---
    document.add_heading('2. Análise de Resultados', level=1)
    
    # Subseção de Participação e Custos
    document.add_heading('Métricas de Participação e Custo', level=2)
    table_metrics = document.add_table(rows=0, cols=2, style='Table Grid')
    table_metrics.columns[0].width = Inches(2.5)
    table_metrics.columns[1].width = Inches(4.0)
    add_info_row(table_metrics, 'Participantes Previstos:', treinamento.participantes_previstos)
    add_info_row(table_metrics, 'Participantes Inscritos:', total_inscritos)
    add_info_row(table_metrics, 'Participantes Presentes:', total_presentes)
    add_info_row(table_metrics, 'Taxa de Presença:', f"{taxa_presenca:.1f}%")
    add_info_row(table_metrics, 'Custo por Participante Efetivo:', f"R$ {custo_por_participante_presente:,.2f}".replace(',', 'X').replace('.', ',').replace('X', '.'))

    document.add_paragraph() # Espaço

    # Subseção de Avaliação e Desempenho
    document.add_heading('Métricas de Avaliação', level=2)
    table_perf = document.add_table(rows=0, cols=2, style='Table Grid')
    table_perf.columns[0].width = Inches(2.5)
    table_perf.columns[1].width = Inches(4.0)
    add_info_row(table_perf, 'Média Geral das Notas:', f"{media_notas:.1f}" if notas else 'N/A')
    add_info_row(table_perf, 'Maior Nota:', f"{maior_nota:.1f}" if notas else 'N/A')
    add_info_row(table_perf, 'Menor Nota:', f"{menor_nota:.1f}" if notas else 'N/A')
    add_info_row(table_perf, 'Taxa de Aprovação (Nota >= 7):', f"{taxa_aprovacao:.1f}%" if notas else 'N/A')


    # --- Seção 3: Gráficos Visuais ---
    document.add_heading('3. Gráficos', level=1)

    # Gráfico 1: Pizza de Presença
    if total_inscritos > 0:
        fig, ax = plt.subplots(figsize=(4, 3))
        labels = ['Presentes', 'Ausentes']
        sizes = [total_presentes, total_ausentes]
        colors = ['#4CAF50', '#F44336']
        ax.pie(sizes, labels=labels, autopct='%1.1f%%', startangle=90, colors=colors, wedgeprops={'edgecolor': 'white'})
        ax.axis('equal') # Garante que a pizza seja um círculo
        ax.set_title('Distribuição de Presença', fontsize=10)

        chart_buffer = io.BytesIO()
        plt.savefig(chart_buffer, format='png', dpi=200, bbox_inches='tight')
        plt.close(fig)
        chart_buffer.seek(0)
        document.add_picture(chart_buffer, width=Inches(3.0))
        # Centralizar a imagem
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Gráfico 2: Barras de Distribuição de Notas
    if notas:
        fig, ax = plt.subplots(figsize=(5, 3.5))
        # Agrupa notas em faixas (0-4, 5-6, 7-8, 9-10)
        bins = [-1, 4.9, 6.9, 8.9, 10.1]
        labels = ['0-4 (Reprovado)', '5-6 (Recuperação)', '7-8 (Bom)', '9-10 (Excelente)']
        nota_counts = pd.cut(notas, bins=bins, labels=labels, right=False).value_counts().sort_index()

        ax.bar(nota_counts.index, nota_counts.values, color=['#F44336', '#FFC107', '#2196F3', '#4CAF50'])
        ax.set_ylabel('Nº de Participantes')
        ax.set_title('Distribuição das Notas de Avaliação', fontsize=10)
        plt.xticks(rotation=15, ha="right") # Rotaciona os labels do eixo x para melhor visualização
        
        chart_buffer = io.BytesIO()
        plt.savefig(chart_buffer, format='png', dpi=200, bbox_inches='tight')
        plt.close(fig)
        chart_buffer.seek(0)

        document.add_paragraph() # Espaço
        document.add_picture(chart_buffer, width=Inches(5.0))
        # Centralizar a imagem
        last_paragraph = document.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    document.add_page_break()

    # --- Seção 4: Lista de Participantes ---
    document.add_heading('4. Lista de Participantes Detalhada', level=1)
    table_participantes = document.add_table(rows=1, cols=4, style='Table Grid')
    hdr_cells = table_participantes.rows[0].cells
    headers = ['Funcionário', 'Matrícula', 'Presença', 'Nota']
    for i, header_text in enumerate(headers):
        hdr_cells[i].text = header_text
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    for p in lista_participantes:
        row_cells = table_participantes.add_row().cells
        row_cells[0].text = p.funcionario.get_full_name() or p.funcionario.username
        row_cells[1].text = p.funcionario.username
        row_cells[2].text = 'Sim' if p.presente else 'Não'
        row_cells[3].text = str(p.nota_avaliacao) if p.nota_avaliacao is not None else 'N/A'
        # Alinhamento das células de status e nota
        row_cells[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER


    # ==================================================================
    # 4. RODAPÉ
    # ==================================================================
    document.add_paragraph()
    p_footer = document.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_footer = p_footer.add_run(f"Relatório gerado em {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')}")
    run_footer.font.size = Pt(8)
    run_footer.font.italic = True


    # ==================================================================
    # 5. SALVAR E RETORNAR O DOCUMENTO
    # ==================================================================
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    return buffer
