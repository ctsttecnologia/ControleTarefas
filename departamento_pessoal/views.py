# departamento_pessoal/views.py
"""
Views do módulo Departamento Pessoal.

Convenções:
- APP_LABEL: constante única do app (usada em decorators e mixins)
- DPBaseMixin: base para todas as CBVs do módulo
- @app_permission_required(APP_LABEL): para FBVs
"""

import io
import logging
from datetime import timedelta
from io import BytesIO

import pandas as pd
from docx import Document as PyDocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl.styles import Alignment, Font, NamedStyle, PatternFill
from openpyxl.worksheet.datavalidation import DataValidation
from weasyprint import HTML

from django.contrib import messages
from django.contrib.messages.views import SuccessMessageMixin
from django.core.exceptions import ObjectDoesNotExist, PermissionDenied
from django.db import IntegrityError, transaction
from django.db.models import Avg, Count, Max, Min, Q
from django.db.models.functions import ExtractMonth, ExtractYear
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.template.loader import render_to_string
from django.urls import reverse, reverse_lazy
from django.utils import timezone
from django.views import View
from django.views.generic import (
    CreateView, DeleteView, DetailView, ListView, TemplateView, UpdateView,
)

from core.decorators import app_permission_required
from core.mixins import (
    AppPermissionMixin, FilialCreateMixin, FuncionarioRequiredMixin,
    ViewFilialScopedMixin,
)
from seguranca_trabalho.models import Funcao

from .forms import (
    AdmissaoForm, CargoForm, DepartamentoForm, DocumentoForm, FuncionarioForm,
    ImportacaoMassaFuncionarioForm, UploadFuncionariosForm,
)
from .models import Cargo, Cliente, Departamento, Documento, Filial, Funcionario
from .services.importacao_massa import gerar_planilha_modelo, processar_planilha

logger = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════════════════════════
# CONSTANTES DO MÓDULO
# ═══════════════════════════════════════════════════════════════════════════════

APP_LABEL = 'departamento_pessoal'
MODULO_NOME = 'Departamento Pessoal'

LIMITE_REGISTROS = 500

COLUNAS_ESPERADAS = [
    'matricula', 'nome_completo', 'data_admissao', 'data_nascimento',
    'email_pessoal', 'telefone', 'salario', 'status', 'sexo',
    'nome_cargo', 'cbo', 'nome_funcao', 'nome_departamento', 'nome_filial', 'nome_cliente',
]


# ═══════════════════════════════════════════════════════════════════════════════
# EXCEÇÕES CUSTOMIZADAS
# ═══════════════════════════════════════════════════════════════════════════════

class ImportacaoError(Exception):
    """Exceção para erros durante a importação da planilha."""

    def __init__(self, message, column_name=None):
        self.message = message
        self.column_name = column_name
        super().__init__(self.message)

    def __str__(self):
        if self.column_name:
            return f"Coluna '{self.column_name}': {self.message}"
        return self.message


# ═══════════════════════════════════════════════════════════════════════════════
# MIXINS — FILIAL E VISIBILIDADE
# ═══════════════════════════════════════════════════════════════════════════════

class FilialAtivaMixin:
    """
    Fornece métodos utilitários para obter a filial ativa do usuário.

    Independente do `ViewFilialScopedMixin` (core), que age no get_queryset().
    Este mixin oferece métodos auxiliares que podem ser usados em qualquer view.
    """

    def get_filial_ativa(self):
        """Retorna a filial ativa do usuário ou None."""
        filial_id = self.request.session.get('active_filial_id')
        if filial_id:
            try:
                return Filial.objects.get(pk=filial_id)
            except Filial.DoesNotExist:
                pass

        user = self.request.user
        filial_ativa = getattr(user, 'filial_ativa', None)
        if filial_ativa:
            return filial_ativa

        try:
            funcionario = Funcionario.objects.select_related('filial').get(usuario=user)
            return funcionario.filial
        except Funcionario.DoesNotExist:
            return None

    def get_filial_ativa_id(self):
        filial = self.get_filial_ativa()
        return filial.id if filial else None

    def filter_queryset_by_filial(self, queryset, filial_field='filial'):
        filial = self.get_filial_ativa()
        if filial:
            return queryset.filter(**{filial_field: filial})
        if not self.request.user.is_superuser:
            return queryset.none()
        return queryset


class DPVisibilityMixin(FilialAtivaMixin):
    """
    Controla visibilidade de registros de DP por perfil de usuário.

    Camadas:
        1. ViewFilialScopedMixin (core)  → escopo de filial
        2. DPVisibilityMixin (este)      → escopo por perfil
        3. Filtros GET nas views         → busca/filtro do usuário

    Regras:
        - Superuser                                          → tudo
        - Permissão `view_all_departamento_pessoal`          → tudo da filial
        - Usuário comum com Funcionario vinculado            → só dados próprios
        - Usuário sem Funcionario                            → nada
    """

    def apply_visibility(self, queryset, funcionario_field=None):
        user = self.request.user

        if user.is_superuser:
            return queryset

        if user.has_perm(f'{APP_LABEL}.view_all_departamento_pessoal'):
            return queryset

        funcionario = getattr(user, 'funcionario', None)
        if not funcionario:
            return queryset.none()

        if funcionario_field:
            return queryset.filter(**{funcionario_field: funcionario})

        if queryset.model == Funcionario:
            return queryset.filter(pk=funcionario.pk)

        # Models auxiliares (Departamento, Cargo) → vê todos da filial
        return queryset


# ═══════════════════════════════════════════════════════════════════════════════
# MIXIN BASE — TODAS AS CBVs DO MÓDULO
# ═══════════════════════════════════════════════════════════════════════════════

class DPBaseMixin(
    FuncionarioRequiredMixin,
    AppPermissionMixin,
    DPVisibilityMixin,
    ViewFilialScopedMixin,
):
    """
    Mixin base para CBVs do Departamento Pessoal.

    MRO (ordem de execução):
        1. FuncionarioRequiredMixin → autenticação + vínculo Funcionario
        2. AppPermissionMixin       → permissão do app
        3. DPVisibilityMixin        → apply_visibility() + FilialAtivaMixin
        4. ViewFilialScopedMixin    → filtra queryset por filial
    """
    app_label_required = APP_LABEL
    modulo_nome = MODULO_NOME


# ═══════════════════════════════════════════════════════════════════════════════
# UPLOAD / IMPORTAÇÃO LEGADA DE FUNCIONÁRIOS
# ═══════════════════════════════════════════════════════════════════════════════

class UploadFuncionariosView(DPBaseMixin, View):
    """Upload e processamento de planilha de funcionários (.xlsx ou .csv)."""

    permission_required = f'{APP_LABEL}.add_funcionario'
    form_class = UploadFuncionariosForm
    template_name = 'departamento_pessoal/upload_funcionarios.html'

    def get(self, request, *args, **kwargs):
        return render(request, self.template_name, {
            'form': self.form_class(),
            'filial_ativa': self.get_filial_ativa(),
        })

    def post(self, request, *args, **kwargs):
        form = self.form_class(request.POST, request.FILES)
        if not form.is_valid():
            return render(request, self.template_name, {'form': form})

        arquivo = request.FILES['arquivo']
        modo_duplicidade = request.POST.get('duplicidade', 'atualizar')

        try:
            df = self._ler_e_validar_planilha(arquivo)
        except ValueError as e:
            messages.error(request, str(e))
            return render(request, self.template_name, {'form': form})

        sucessos, linhas_erro = self._processar_dataframe(df, modo_duplicidade)

        if linhas_erro:
            request.session['upload_erros'] = linhas_erro
            messages.error(
                request,
                f"A importação falhou. {len(linhas_erro)} erro(s) encontrado(s). "
                "Nenhum funcionário foi salvo. Baixe o relatório de erros."
            )
        else:
            messages.success(
                request,
                f"Importação concluída! {sucessos} funcionário(s) criado(s)/atualizado(s)."
            )

        return redirect(reverse('departamento_pessoal:upload_funcionarios'))

    def _ler_e_validar_planilha(self, arquivo):
        nome = arquivo.name.lower()

        try:
            if nome.endswith('.csv'):
                conteudo = arquivo.read()
                arquivo.seek(0)

                for encoding in ('utf-8', 'latin-1', 'cp1252'):
                    try:
                        df = pd.read_csv(
                            io.BytesIO(conteudo),
                            dtype=str,
                            encoding=encoding,
                            sep=None,
                            engine='python',
                        ).fillna('')
                        break
                    except (UnicodeDecodeError, pd.errors.ParserError):
                        continue
                else:
                    raise ValueError(
                        'Não foi possível ler o arquivo CSV. '
                        'Verifique a codificação (recomendado: UTF-8).'
                    )

            elif nome.endswith('.xlsx'):
                df = pd.read_excel(arquivo, dtype=str).fillna('')

            else:
                raise ValueError('Formato não suportado. Use .xlsx ou .csv')

        except ValueError:
            raise
        except Exception as e:
            logger.exception("Erro ao ler planilha de upload")
            raise ValueError(f'Erro ao ler o arquivo: {e}')

        df.columns = df.columns.str.strip().str.lower().str.replace(' ', '_')

        faltantes = set(COLUNAS_ESPERADAS) - set(df.columns)
        if faltantes:
            raise ValueError(
                f"Colunas obrigatórias ausentes: {', '.join(sorted(faltantes))}"
            )

        df = df.loc[~(df == '').all(axis=1)]

        if len(df) == 0:
            raise ValueError('A planilha está vazia.')

        if len(df) > LIMITE_REGISTROS:
            raise ValueError(
                f'Planilha com {len(df)} registros excede o limite de {LIMITE_REGISTROS}.'
            )

        return df

    def _processar_dataframe(self, df, modo_duplicidade='atualizar'):
        linhas_erro = []
        sucessos = 0

        try:
            with transaction.atomic():
                for index, row in df.iterrows():
                    linha_num = index + 2
                    try:
                        self._processar_linha(row, modo_duplicidade)
                        sucessos += 1
                    except ImportacaoError as e:
                        linha_erro = row.to_dict()
                        linha_erro['Erro'] = f"Linha {linha_num}, {e}"
                        linhas_erro.append(linha_erro)
                    except Exception as e:
                        logger.exception(f"Erro inesperado na linha {linha_num}")
                        linha_erro = row.to_dict()
                        linha_erro['Erro'] = f"Linha {linha_num}: Erro inesperado — {e}"
                        linhas_erro.append(linha_erro)

                if linhas_erro:
                    raise IntegrityError("Erros encontrados durante a importação.")

        except IntegrityError:
            return 0, linhas_erro

        return sucessos, linhas_erro

    def _processar_linha(self, row, modo_duplicidade='atualizar'):
        # --- Validação de campos obrigatórios ---
        matricula = row.get('matricula', '').strip()
        if not matricula:
            raise ImportacaoError("Este campo é obrigatório.", column_name='matricula')

        nome_completo = row.get('nome_completo', '').strip()
        if not nome_completo:
            raise ImportacaoError("Este campo é obrigatório.", column_name='nome_completo')

        if modo_duplicidade == 'pular' and Funcionario.objects.filter(matricula=matricula).exists():
            raise ImportacaoError(
                f"Matrícula '{matricula}' já existe. Linha pulada conforme configuração.",
                column_name='matricula',
            )

        # --- Datas ---
        data_admissao = self._parse_data(row, 'data_admissao', obrigatorio=True)
        data_nascimento = self._parse_data(row, 'data_nascimento', obrigatorio=False)

        # --- Filial (obrigatória) ---
        nome_filial = row.get('nome_filial', '').strip()
        if not nome_filial:
            raise ImportacaoError("Este campo é obrigatório.", column_name='nome_filial')
        try:
            filial = Filial.objects.get(nome__iexact=nome_filial)
        except Filial.DoesNotExist:
            raise ImportacaoError(
                f"Filial '{nome_filial}' não encontrada. Cadastre-a primeiro.",
                column_name='nome_filial',
            )

        # --- Cargo ---
        nome_cargo = row.get('nome_cargo', '').strip()
        if not nome_cargo:
            raise ImportacaoError("Este campo é obrigatório.", column_name='nome_cargo')
        cbo = row.get('cbo', '').strip()
        cargo, _ = Cargo.objects.get_or_create(
            nome__iexact=nome_cargo,
            filial=filial,
            defaults={'nome': nome_cargo.title(), 'filial': filial, 'cbo': cbo},
        )

        # --- Função (opcional) ---
        funcao = None
        nome_funcao = row.get('nome_funcao', '').strip()
        if nome_funcao:
            funcao, _ = Funcao.objects.get_or_create(
                nome__iexact=nome_funcao,
                filial=filial,
                defaults={'nome': nome_funcao.title(), 'filial': filial},
            )

        # --- Departamento ---
        nome_dpto = row.get('nome_departamento', '').strip()
        if not nome_dpto:
            raise ImportacaoError("Este campo é obrigatório.", column_name='nome_departamento')
        departamento, _ = Departamento.objects.get_or_create(
            nome__iexact=nome_dpto,
            filial=filial,
            defaults={'nome': nome_dpto.title(), 'filial': filial},
        )

        # --- Cliente (opcional) ---
        cliente = None
        nome_cliente = row.get('nome_cliente', '').strip()
        if nome_cliente:
            try:
                cliente = Cliente.objects.get(nome__iexact=nome_cliente)
            except Cliente.DoesNotExist:
                raise ImportacaoError(
                    f"Cliente '{nome_cliente}' não encontrado. Cadastre-o primeiro.",
                    column_name='nome_cliente',
                )

        # --- Status ---
        status = row.get('status', '').upper().strip()
        if not status:
            raise ImportacaoError("Este campo é obrigatório.", column_name='status')
        if status not in dict(Funcionario.STATUS_CHOICES):
            opcoes = ', '.join(dict(Funcionario.STATUS_CHOICES).keys())
            raise ImportacaoError(
                f"Valor '{row.get('status')}' inválido. Opções: {opcoes}",
                column_name='status',
            )

        # --- Sexo (opcional) ---
        sexo = row.get('sexo', '').upper().strip() or None
        if sexo and sexo not in dict(Funcionario.SEXO_CHOICES):
            opcoes = ', '.join(dict(Funcionario.SEXO_CHOICES).keys())
            raise ImportacaoError(
                f"Valor '{row.get('sexo')}' inválido. Opções: {opcoes}",
                column_name='sexo',
            )

        # --- Salário ---
        salario_str = row.get('salario', '').strip()
        salario = 0.00
        if salario_str:
            try:
                salario = float(salario_str.replace(',', '.'))
            except ValueError:
                raise ImportacaoError(
                    f"Valor '{salario_str}' não é número válido.",
                    column_name='salario',
                )

        # --- Email (validação de duplicidade) ---
        email = row.get('email_pessoal', '').strip() or None
        if email and Funcionario.objects.filter(
            email_pessoal__iexact=email
        ).exclude(matricula=matricula).exists():
            raise ImportacaoError(
                f"E-mail '{email}' já cadastrado para outro funcionário.",
                column_name='email_pessoal',
            )

        # --- Persistência ---
        Funcionario.objects.update_or_create(
            matricula=matricula,
            defaults={
                'nome_completo': nome_completo.upper(),
                'data_admissao': data_admissao.date(),
                'data_nascimento': data_nascimento.date() if data_nascimento is not None else None,
                'cargo': cargo,
                'funcao': funcao,
                'departamento': departamento,
                'filial': filial,
                'cliente': cliente,
                'status': status,
                'salario': salario,
                'email_pessoal': email,
                'telefone': row.get('telefone', '').strip(),
                'sexo': sexo,
            },
        )

    @staticmethod
    def _parse_data(row, campo, obrigatorio=False):
        """Converte string de data para datetime; retorna None se opcional e vazia."""
        valor = row.get(campo, '').strip()
        if not valor:
            if obrigatorio:
                raise ImportacaoError("Este campo é obrigatório.", column_name=campo)
            return None
        data = pd.to_datetime(valor, dayfirst=True, errors='coerce')
        if pd.isna(data):
            raise ImportacaoError(
                f"Valor '{valor}' em formato inválido. Use DD-MM-AAAA.",
                column_name=campo,
            )
        return data


# ═══════════════════════════════════════════════════════════════════════════════
# FBVs — DOWNLOAD DE TEMPLATE E RELATÓRIO DE ERROS
# ═══════════════════════════════════════════════════════════════════════════════

@app_permission_required(APP_LABEL)
def baixar_modelo_funcionarios(request):
    """Gera planilha modelo .xlsx com formatação profissional."""
    df_modelo = pd.DataFrame(columns=COLUNAS_ESPERADAS)
    output = BytesIO()

    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df_modelo.to_excel(writer, index=False, sheet_name='Funcionarios')

        workbook = writer.book
        worksheet = writer.sheets['Funcionarios']

        text_style = NamedStyle(name='text_style', number_format='@')
        workbook.add_named_style(text_style)

        header_font = Font(bold=True, color='FFFFFF', name='Calibri')
        header_fill = PatternFill(start_color='4F81BD', end_color='4F81BD', fill_type='solid')
        header_alignment = Alignment(horizontal='center', vertical='center')

        for col_num, column_title in enumerate(df_modelo.columns, 1):
            cell = worksheet.cell(row=1, column=col_num)
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment

            if 'data_' in column_title:
                for r in range(2, 5000):
                    worksheet.cell(row=r, column=col_num).style = text_style

            worksheet.column_dimensions[cell.column_letter].width = len(column_title) + 5

        # Validações de lista
        status_options = f'"{",".join(c[0] for c in Funcionario.STATUS_CHOICES)}"'
        dv_status = DataValidation(type='list', formula1=status_options, allow_blank=True)
        dv_status.add('H2:H1048576')
        worksheet.add_data_validation(dv_status)

        sexo_options = f'"{",".join(c[0] for c in Funcionario.SEXO_CHOICES)}"'
        dv_sexo = DataValidation(type='list', formula1=sexo_options, allow_blank=True)
        dv_sexo.add('I2:I1048576')
        worksheet.add_data_validation(dv_sexo)

        worksheet.freeze_panes = 'A2'

    output.seek(0)
    response = HttpResponse(
        output,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = 'attachment; filename="modelo_importacao_funcionarios.xlsx"'
    return response


@app_permission_required(APP_LABEL)
def baixar_relatorio_erros(request):
    """Gera Excel com as linhas que continham erros durante o último upload."""
    linhas_erro = request.session.pop('upload_erros', [])

    if not linhas_erro:
        messages.warning(request, "Não há relatório de erros para baixar.")
        return redirect(reverse('departamento_pessoal:upload_funcionarios'))

    df = pd.DataFrame(linhas_erro)
    output = BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Erros')
    output.seek(0)

    response = HttpResponse(
        output,
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = 'attachment; filename="relatorio_erros_importacao.xlsx"'
    return response


# ═══════════════════════════════════════════════════════════════════════════════
# CRUD — FUNCIONÁRIOS
# ═══════════════════════════════════════════════════════════════════════════════

class FuncionarioListView(DPBaseMixin, ListView):
    permission_required = f'{APP_LABEL}.view_funcionario'
    model = Funcionario
    template_name = 'departamento_pessoal/lista_funcionarios.html'
    context_object_name = 'funcionarios'
    paginate_by = 15

    def dispatch(self, request, *args, **kwargs):
        """Bloqueia acesso se usuário não tem vínculo de funcionário."""
        if not request.user.is_superuser:
            try:
                _ = request.user.funcionario
            except ObjectDoesNotExist:
                return render(request, 'departamento_pessoal/acesso_negado.html', {
                    'titulo': 'Acesso Restrito',
                    'mensagem': (
                        'Sua conta não está vinculada a um registro de funcionário, '
                        'por isso não é possível acessar o módulo de Departamento Pessoal.'
                    ),
                }, status=403)
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        queryset = super().get_queryset().select_related(
            'cargo', 'departamento'
        ).order_by('nome_completo')

        queryset = self.apply_visibility(queryset)

        query = self.request.GET.get('q')
        if query:
            queryset = queryset.filter(
                Q(nome_completo__icontains=query) |
                Q(matricula__icontains=query) |
                Q(cargo__nome__icontains=query)
            )
        return queryset

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['filial_ativa'] = self.get_filial_ativa()
        context['current_q'] = self.request.GET.get('q', '')
        return context


class FuncionarioCreateView(DPBaseMixin, FilialCreateMixin, CreateView):
    permission_required = f'{APP_LABEL}.add_funcionario'
    model = Funcionario
    form_class = FuncionarioForm
    template_name = 'departamento_pessoal/funcionario_form.html'
    extra_context = {'titulo_pagina': 'Cadastrar Novo Funcionário'}

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['request'] = self.request
        return kwargs

    def get_success_url(self):
        return reverse_lazy(
            'departamento_pessoal:detalhe_funcionario',
            kwargs={'pk': self.object.pk},
        )


class FuncionarioDetailView(DPBaseMixin, DetailView):
    permission_required = f'{APP_LABEL}.view_funcionario'
    model = Funcionario
    template_name = 'departamento_pessoal/detalhe_funcionario.html'
    context_object_name = 'funcionario'

    def get_queryset(self):
        qs = super().get_queryset().select_related('usuario', 'cargo', 'departamento')
        return self.apply_visibility(qs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['filial_ativa'] = self.get_filial_ativa()
        return context


class FuncionarioUpdateView(DPBaseMixin, SuccessMessageMixin, UpdateView):
    permission_required = f'{APP_LABEL}.change_funcionario'
    model = Funcionario
    form_class = FuncionarioForm
    template_name = 'departamento_pessoal/funcionario_form.html'
    success_message = '🔄 Dados do funcionário atualizados com sucesso!'

    def get_queryset(self):
        qs = super().get_queryset().select_related('usuario', 'cargo', 'departamento')
        return self.apply_visibility(qs)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['request'] = self.request
        return kwargs

    def get_success_url(self):
        return reverse_lazy(
            'departamento_pessoal:detalhe_funcionario',
            kwargs={'pk': self.object.pk},
        )

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo_pagina'] = f"Editar: {self.object.nome_completo}"
        context['filial_ativa'] = self.get_filial_ativa()
        return context


class FuncionarioDeleteView(DPBaseMixin, DeleteView):
    """
    Exclusão/inativação de funcionário.
    Suporta ações via POST: 'inativar' ou 'excluir'.
    """
    permission_required = f'{APP_LABEL}.delete_funcionario'
    model = Funcionario
    template_name = 'departamento_pessoal/confirm_delete.html'
    context_object_name = 'funcionario'
    success_url = reverse_lazy('departamento_pessoal:lista_funcionarios')

    def get_queryset(self):
        return self.apply_visibility(super().get_queryset())

    def form_valid(self, form):
        funcionario = self.get_object()
        action = self.request.POST.get('action')

        if action == 'inativar':
            funcionario.status = 'INATIVO'
            funcionario.save(update_fields=['status'])
            messages.warning(
                self.request,
                f"⚠️ O funcionário '{funcionario.nome_completo}' foi INATIVADO."
            )
            return redirect(self.success_url)

        nome_completo = funcionario.nome_completo
        response = super().form_valid(form)
        messages.error(
            self.request,
            f"🗑️ O funcionário '{nome_completo}' foi EXCLUÍDO PERMANENTEMENTE."
        )
        return response


# ═══════════════════════════════════════════════════════════════════════════════
# ADMISSÃO
# ═══════════════════════════════════════════════════════════════════════════════

class FuncionarioAdmissaoView(DPBaseMixin, SuccessMessageMixin, UpdateView):
    permission_required = f'{APP_LABEL}.change_funcionario'
    model = Funcionario
    form_class = AdmissaoForm
    template_name = 'departamento_pessoal/admissao_form.html'

    def get_queryset(self):
        return self.apply_visibility(super().get_queryset())

    def get_success_message(self, cleaned_data):
        return f"✅ Dados de admissão de '{self.object.nome_completo}' salvos com sucesso!"

    def get_success_url(self):
        return reverse(
            'departamento_pessoal:detalhe_funcionario',
            kwargs={'pk': self.object.pk},
        )

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        prefixo = 'Editar Admissão de' if self.object.data_admissao else 'Registrar Admissão para'
        context['titulo_pagina'] = f"{prefixo} {self.object.nome_completo}"
        context['filial_ativa'] = self.get_filial_ativa()
        return context


# ═══════════════════════════════════════════════════════════════════════════════
# CRUD — DEPARTAMENTOS
# ═══════════════════════════════════════════════════════════════════════════════

class DepartamentoListView(DPBaseMixin, ListView):
    permission_required = f'{APP_LABEL}.view_departamento'
    model = Departamento
    template_name = 'departamento_pessoal/lista_departamento.html'
    context_object_name = 'departamentos'

    def get_queryset(self):
        qs = super().get_queryset().filter(ativo=True)
        return self.apply_visibility(qs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['filial_ativa'] = self.get_filial_ativa()
        return context


class DepartamentoCreateView(DPBaseMixin, FilialCreateMixin, SuccessMessageMixin, CreateView):
    permission_required = f'{APP_LABEL}.add_departamento'
    model = Departamento
    form_class = DepartamentoForm
    template_name = 'departamento_pessoal/departamento_form.html'
    success_url = reverse_lazy('departamento_pessoal:lista_departamento')
    success_message = '✅ Departamento criado com sucesso!'
    extra_context = {'titulo_pagina': 'Novo Departamento'}


class DepartamentoUpdateView(DPBaseMixin, SuccessMessageMixin, UpdateView):
    permission_required = f'{APP_LABEL}.change_departamento'
    model = Departamento
    form_class = DepartamentoForm
    template_name = 'departamento_pessoal/departamento_form.html'
    success_url = reverse_lazy('departamento_pessoal:lista_departamento')
    success_message = '🔄 Departamento atualizado com sucesso!'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo_pagina'] = f"Editar Departamento: {self.object.nome}"
        return context


# ═══════════════════════════════════════════════════════════════════════════════
# CRUD — CARGOS
# ═══════════════════════════════════════════════════════════════════════════════

class CargoListView(DPBaseMixin, ListView):
    permission_required = f'{APP_LABEL}.view_cargo'
    model = Cargo
    template_name = 'departamento_pessoal/lista_cargo.html'
    context_object_name = 'cargos'
    paginate_by = 20

    def get_queryset(self):
        qs = super().get_queryset().filter(ativo=True)
        return self.apply_visibility(qs)

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['filial_ativa'] = self.get_filial_ativa()
        return context


class CargoCreateView(DPBaseMixin, FilialCreateMixin, SuccessMessageMixin, CreateView):
    permission_required = f'{APP_LABEL}.add_cargo'
    model = Cargo
    form_class = CargoForm
    template_name = 'departamento_pessoal/cargo_form.html'
    success_url = reverse_lazy('departamento_pessoal:lista_cargo')
    success_message = '✅ Cargo criado com sucesso!'
    extra_context = {'titulo_pagina': 'Novo Cargo'}


class CargoUpdateView(DPBaseMixin, SuccessMessageMixin, UpdateView):
    permission_required = f'{APP_LABEL}.change_cargo'
    model = Cargo
    form_class = CargoForm
    template_name = 'departamento_pessoal/cargo_form.html'
    success_url = reverse_lazy('departamento_pessoal:lista_cargo')
    success_message = '🔄 Cargo atualizado com sucesso!'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['titulo_pagina'] = f"Editar Cargo: {self.object.nome}"
        return context


# ═══════════════════════════════════════════════════════════════════════════════
# CRUD — DOCUMENTOS
# ═══════════════════════════════════════════════════════════════════════════════

class _DocumentoFilialScopedMixin(FilialAtivaMixin):
    """Garante que documentos só sejam acessados dentro da filial ativa."""

    def get_queryset(self):
        qs = super().get_queryset()
        filial = self.get_filial_ativa()

        if not filial and not self.request.user.is_superuser:
            raise PermissionDenied("Nenhuma filial selecionada.")

        if filial:
            qs = qs.filter(funcionario__filial=filial)
        return qs


class DocumentoListView(DPBaseMixin, _DocumentoFilialScopedMixin, ListView):
    permission_required = f'{APP_LABEL}.view_documento'
    model = Documento
    template_name = 'departamento_pessoal/lista_documentos.html'
    context_object_name = 'documentos'
    paginate_by = 20

    def get_queryset(self):
        filial = self.get_filial_ativa()
        if not filial:
            messages.error(self.request, "Selecione uma filial para ver os documentos.")
            return self.model.objects.none()

        queryset = Documento.objects.filter(
            funcionario__filial=filial
        ).select_related('funcionario', 'funcionario__cargo')

        tipo = self.request.GET.get('tipo', '')
        if tipo:
            queryset = queryset.filter(tipo_documento=tipo)

        query = self.request.GET.get('q', '')
        if query:
            queryset = queryset.filter(
                Q(funcionario__nome_completo__icontains=query) |
                Q(numero__icontains=query)
            )

        return queryset.order_by('funcionario__nome_completo', 'tipo_documento')

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['tipos_documento'] = Documento.TIPO_CHOICES
        context['filial_ativa'] = self.get_filial_ativa()
        context['current_q'] = self.request.GET.get('q', '')
        context['current_tipo'] = self.request.GET.get('tipo', '')
        return context


class DocumentoCreateView(DPBaseMixin, FilialCreateMixin, CreateView):
    permission_required = f'{APP_LABEL}.add_documento'
    model = Documento
    form_class = DocumentoForm
    template_name = 'departamento_pessoal/documento_form.html'

    def get_initial(self):
        initial = super().get_initial()
        funcionario_pk = self.kwargs.get('funcionario_pk')
        if funcionario_pk:
            qs = Funcionario.objects.filter(pk=funcionario_pk)
            filial = self.get_filial_ativa()
            if filial:
                qs = qs.filter(filial=filial)
            funcionario = qs.first()
            if funcionario:
                initial['funcionario'] = funcionario
        return initial

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        funcionario_pk = self.kwargs.get('funcionario_pk')
        if funcionario_pk:
            context['funcionario'] = Funcionario.objects.filter(pk=funcionario_pk).first()
        context['filial_ativa'] = self.get_filial_ativa()
        return context

    def form_valid(self, form):
        try:
            messages.success(self.request, '✅ Documento adicionado com sucesso.')
            return super().form_valid(form)
        except IntegrityError:
            form.add_error(
                None,
                'Este funcionário já possui um documento deste tipo com o mesmo número.',
            )
            return self.form_invalid(form)

    def get_success_url(self):
        return reverse(
            'departamento_pessoal:detalhe_funcionario',
            kwargs={'pk': self.object.funcionario.pk},
        )


class DocumentoUpdateView(DPBaseMixin, _DocumentoFilialScopedMixin, UpdateView):
    permission_required = f'{APP_LABEL}.change_documento'
    model = Documento
    form_class = DocumentoForm
    template_name = 'departamento_pessoal/documento_form.html'

    def get_success_url(self):
        messages.success(self.request, '🔄 Documento atualizado com sucesso.')
        return reverse(
            'departamento_pessoal:detalhe_funcionario',
            kwargs={'pk': self.object.funcionario.pk},
        )

    def form_valid(self, form):
        try:
            return super().form_valid(form)
        except IntegrityError:
            form.add_error(
                'tipo_documento',
                'Já existe outro registro deste tipo de documento para este funcionário.',
            )
            return self.form_invalid(form)


class DocumentoDeleteView(DPBaseMixin, _DocumentoFilialScopedMixin, DeleteView):
    permission_required = f'{APP_LABEL}.delete_documento'
    model = Documento
    template_name = 'departamento_pessoal/documento_confirm_delete.html'
    context_object_name = 'documento'

    def get_success_url(self):
        messages.success(self.request, '🗑️ Documento excluído com sucesso.')
        return reverse_lazy(
            'departamento_pessoal:detalhe_funcionario',
            kwargs={'pk': self.object.funcionario.pk},
        )


# ═══════════════════════════════════════════════════════════════════════════════
# PAINEL DP — DASHBOARD
# ═══════════════════════════════════════════════════════════════════════════════

class PainelDPView(DPBaseMixin, TemplateView):
    permission_required = f'{APP_LABEL}.view_funcionario'
    template_name = 'departamento_pessoal/painel_dp.html'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        filial = self.get_filial_ativa()
        if not filial and not self.request.user.is_superuser:
            context['permission_denied'] = True
            messages.error(self.request, "Selecione uma filial para visualizar o painel.")
            return context

        # --- Querysets base ---
        if filial:
            funcionarios_qs = Funcionario.objects.filter(filial=filial)
            departamentos_qs = Departamento.objects.filter(filial=filial)
            cargos_count = Cargo.objects.filter(filial=filial, ativo=True).count()
        else:
            funcionarios_qs = Funcionario.objects.all()
            departamentos_qs = Departamento.objects.all()
            cargos_count = Cargo.objects.filter(ativo=True).count()

        context['filial_ativa'] = filial
        context['ultimos_funcionarios'] = funcionarios_qs.select_related(
            'cargo', 'departamento'
        ).order_by('-data_admissao')[:5]

        # --- KPIs principais ---
        context['total_funcionarios_ativos'] = funcionarios_qs.filter(status='ATIVO').count()
        context['total_departamentos'] = departamentos_qs.filter(ativo=True).count()
        context['total_cargos'] = cargos_count

        # --- 1. Funcionários por Departamento ---
        func_por_depto = departamentos_qs.filter(ativo=True).annotate(
            num_funcionarios=Count(
                'funcionarios', filter=Q(funcionarios__status='ATIVO')
            )
        ).values('nome', 'num_funcionarios').order_by('-num_funcionarios')

        context['depto_labels'] = [d['nome'] for d in func_por_depto]
        context['depto_data'] = [d['num_funcionarios'] for d in func_por_depto]

        # --- 2. Distribuição por Status ---
        dist_status = funcionarios_qs.values('status').annotate(
            total=Count('status')
        ).order_by('status')
        status_map = dict(Funcionario.STATUS_CHOICES)
        context['status_labels'] = [status_map.get(s['status'], s['status']) for s in dist_status]
        context['status_data'] = [s['total'] for s in dist_status]

        # --- 3. Salário Médio por Cargo (TOP 10) ---
        salarios_por_cargo = funcionarios_qs.filter(
            status='ATIVO', salario__gt=0
        ).values('cargo__nome').annotate(
            salario_medio=Avg('salario'),
            total_funcionarios=Count('id'),
        ).order_by('-salario_medio')[:10]

        context['cargo_salario_labels'] = [item['cargo__nome'] for item in salarios_por_cargo]
        context['cargo_salario_data'] = [float(item['salario_medio']) for item in salarios_por_cargo]

        # --- 4. Admissões últimos 12 meses ---
        doze_meses_atras = timezone.now().date() - timedelta(days=365)
        admissoes_por_mes = funcionarios_qs.filter(
            data_admissao__gte=doze_meses_atras
        ).annotate(
            mes=ExtractMonth('data_admissao'),
            ano=ExtractYear('data_admissao'),
        ).values('ano', 'mes').annotate(total=Count('id')).order_by('ano', 'mes')

        context['admissoes_meses_labels'] = [f"{a['mes']}/{a['ano']}" for a in admissoes_por_mes]
        context['admissoes_meses_data'] = [a['total'] for a in admissoes_por_mes]

        # --- 5. Dispersão Idade vs Salário ---
        hoje = timezone.now().date()
        funcionarios_idade_salario = funcionarios_qs.filter(
            status='ATIVO', data_nascimento__isnull=False, salario__gt=0,
        ).values('data_nascimento', 'salario', 'sexo')[:50]

        context['dispersao_data'] = [
            {
                'x': (hoje - f['data_nascimento']).days // 365,
                'y': float(f['salario']),
                'sexo': f['sexo'] or 'N/A',
            }
            for f in funcionarios_idade_salario if f['data_nascimento']
        ]

        # --- 6. Estatísticas Detalhadas ---
        salarios_ativos = funcionarios_qs.filter(status='ATIVO', salario__gt=0)
        if salarios_ativos.exists():
            agg = salarios_ativos.aggregate(
                medio=Avg('salario'), maximo=Max('salario'), minimo=Min('salario'),
            )
            context['salario_medio'] = agg['medio']
            context['salario_maximo'] = agg['maximo']
            context['salario_minimo'] = agg['minimo']

            idades = [
                hoje.year - f.data_nascimento.year - (
                    (hoje.month, hoje.day) < (f.data_nascimento.month, f.data_nascimento.day)
                )
                for f in funcionarios_qs.filter(status='ATIVO', data_nascimento__isnull=False)
            ]
            context['idade_media'] = sum(idades) / len(idades) if idades else 0
        else:
            context['salario_medio'] = 0
            context['salario_maximo'] = 0
            context['salario_minimo'] = 0
            context['idade_media'] = 0

        # --- 7. Distribuição por Gênero ---
        dist_genero = funcionarios_qs.filter(status='ATIVO').values('sexo').annotate(
            total=Count('id')
        )
        genero_map = dict(Funcionario.SEXO_CHOICES)
        context['genero_labels'] = [
            genero_map.get(g['sexo'], g['sexo'] or 'Não informado') for g in dist_genero
        ]
        context['genero_data'] = [g['total'] for g in dist_genero]

        context['titulo_pagina'] = 'Painel de Controle DP - Analytics'
        return context


# ═══════════════════════════════════════════════════════════════════════════════
# EXPORTAÇÕES
# ═══════════════════════════════════════════════════════════════════════════════

class _BaseExportView(DPBaseMixin, View):
    permission_required = f'{APP_LABEL}.view_funcionario'

    def get_scoped_queryset(self):
        filial = self.get_filial_ativa()
        if not filial and not self.request.user.is_superuser:
            raise PermissionDenied("Nenhuma filial selecionada para exportar dados.")

        qs = Funcionario.objects.select_related('cargo', 'departamento', 'filial')
        if filial:
            qs = qs.filter(filial=filial)
        return qs


class ExportarFuncionariosExcelView(_BaseExportView):
    def get(self, request, *args, **kwargs):
        funcionarios = self.get_scoped_queryset().all()

        data = [
            {
                'Matrícula': f.matricula,
                'Nome Completo': f.nome_completo,
                'Cargo': f.cargo.nome if f.cargo else '-',
                'Departamento': f.departamento.nome if f.departamento else '-',
                'Data de Admissão': f.data_admissao.strftime('%d/%m/%Y') if f.data_admissao else '-',
                'Status': f.get_status_display(),
            }
            for f in funcionarios
        ]
        df = pd.DataFrame(data)

        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = 'attachment; filename="relatorio_funcionarios.xlsx"'
        df.to_excel(response, index=False)
        return response


class ExportarFuncionariosPDFView(_BaseExportView):
    def get(self, request, *args, **kwargs):
        funcionarios = self.get_scoped_queryset().filter(status='ATIVO')
        context = {
            'funcionarios': funcionarios,
            'data_emissao': timezone.now().strftime('%d/%m/%Y às %H:%M'),
            'filial_ativa': self.get_filial_ativa(),
        }
        html_string = render_to_string(
            'departamento_pessoal/relatorio_funcionarios_pdf.html', context
        )
        pdf_file = HTML(
            string=html_string, base_url=request.build_absolute_uri()
        ).write_pdf()

        response = HttpResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="relatorio_funcionarios.pdf"'
        return response


class ExportarFuncionariosWordView(_BaseExportView):
    def get(self, request, *args, **kwargs):
        funcionarios = self.get_scoped_queryset().filter(status='ATIVO')

        document = PyDocxDocument()
        document.add_heading('Relatório de Colaboradores Ativos', level=1)

        nome_filial = funcionarios.first().filial.nome if funcionarios.exists() else 'N/A'
        document.add_paragraph(f"Filial: {nome_filial}")
        document.add_paragraph(f"Data de Emissão: {timezone.now().strftime('%d/%m/%Y às %H:%M')}")
        document.add_paragraph(f"Total de Colaboradores Listados: {len(funcionarios)}")
        document.add_paragraph()

        colunas = ['Matrícula', 'Nome Completo', 'Cargo', 'Departamento', 'Data de Admissão']
        tabela = document.add_table(rows=1, cols=len(colunas))
        tabela.style = 'Table Grid'

        for i, nome_coluna in enumerate(colunas):
            cell = tabela.rows[0].cells[i]
            cell.text = nome_coluna
            cell.paragraphs[0].runs[0].font.bold = True
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

        for f in funcionarios:
            row = tabela.add_row().cells
            row[0].text = f.matricula or '-'
            row[1].text = f.nome_completo
            row[2].text = f.cargo.nome if f.cargo else '-'
            row[3].text = f.departamento.nome if f.departamento else '-'
            row[4].text = f.data_admissao.strftime('%d/%m/%Y') if f.data_admissao else '-'

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        )
        response['Content-Disposition'] = 'attachment; filename="relatorio_colaboradores_ativos.docx"'
        return response


# ═══════════════════════════════════════════════════════════════════════════════
# IMPORTAÇÃO EM MASSA (serviço novo)
# ═══════════════════════════════════════════════════════════════════════════════

@app_permission_required(APP_LABEL)
def download_modelo_funcionarios_view(request):
    """Gera e retorna planilha modelo para download."""
    filial = getattr(request.user, 'filial_ativa', None)
    buffer = gerar_planilha_modelo(filial)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
    )
    response['Content-Disposition'] = 'attachment; filename="modelo_importacao_funcionarios.xlsx"'
    return response


@app_permission_required(APP_LABEL)
def importacao_massa_funcionarios_view(request):
    """Upload e processamento da planilha de funcionários."""
    if not request.user.has_perm(f'{APP_LABEL}.add_funcionario'):
        raise PermissionDenied("Você não tem permissão para importar funcionários.")

    if request.method == 'POST':
        form = ImportacaoMassaFuncionarioForm(request.POST, request.FILES)
        if form.is_valid():
            arquivo = form.cleaned_data['arquivo']
            filial = getattr(request.user, 'filial_ativa', None)

            if not filial and not request.user.is_superuser:
                messages.error(
                    request,
                    "Nenhuma filial ativa. Selecione uma filial no menu superior."
                )
                return redirect(request.path)

            resultado = processar_planilha(arquivo, filial)

            return render(
                request,
                'departamento_pessoal/importacao_massa_resultado.html',
                {'resultado': resultado, 'form': ImportacaoMassaFuncionarioForm()},
            )
    else:
        form = ImportacaoMassaFuncionarioForm()

    return render(
        request,
        'departamento_pessoal/importacao_massa.html',
        {'form': form},
    )

