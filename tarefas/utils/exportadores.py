
"""
Exportadores de relatórios — PDF, DOCX e XLSX com identidade visual unificada.
"""
from io import BytesIO

from django.template.loader import render_to_string
from django.http import HttpResponse
from django.utils import timezone

from weasyprint import HTML
from docx import Document
from docx.shared import Cm

from .docx_styles import (
    aplicar_cabecalho_word,
    adicionar_titulo_secao,
    adicionar_kpis,
    aplicar_estilo_tabela_word,
    adicionar_badge_cell,
)


# ============================================================
# PDF
# ============================================================
def gerar_pdf_relatorio(context, request=None):
    """Gera PDF do relatório de tarefas com identidade visual padronizada."""
    context.setdefault('now', timezone.now())
    if request is not None:
        context.setdefault('request', request)

    html_string = render_to_string(
        'reports/relatorio_tarefas.html',   # ✅ path correto
        context
    )

    base_url = None
    req = context.get('request')
    if req is not None:
        base_url = req.build_absolute_uri('/')

    pdf_file = HTML(string=html_string, base_url=base_url).write_pdf()

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = (
        f'attachment; filename="relatorio_tarefas_{timezone.now():%Y%m%d_%H%M}.pdf"'
    )
    return response


# ============================================================
# DOCX
# ============================================================
def gerar_docx_relatorio(context):
    """Gera DOCX do relatório de tarefas com identidade visual padronizada."""
    doc = Document()

    # ===== Margens =====
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.2)
        section.right_margin = Cm(1.2)

    # ===== CABEÇALHO =====
    aplicar_cabecalho_word(
        doc,
        titulo="Relatório de Tarefas",
        subtitulo="Acompanhamento e análise de tarefas do sistema",
        data_emissao=timezone.now().strftime("%d/%m/%Y %H:%M"),
    )

    # ===== KPIs =====
    adicionar_kpis(doc, [
        {"label": "Total",        "value": context.get('total_tarefas', 0),      "tipo": "info"},
        {"label": "Concluídas",   "value": context.get('tarefas_concluidas', 0), "tipo": "success"},
        {"label": "Em Andamento", "value": context.get('tarefas_andamento', 0),  "tipo": "warning"},
        {"label": "Atrasadas",    "value": context.get('tarefas_atrasadas', 0),  "tipo": "danger"},
    ])
    doc.add_paragraph()

    # ===== ANÁLISE POR STATUS =====
    analise_status = context.get('analise_status', [])
    if analise_status:
        adicionar_titulo_secao(doc, "Análise por Status")
        _montar_tabela_analise(doc, "Status", analise_status)
        doc.add_paragraph()

    # ===== ANÁLISE POR PRIORIDADE =====
    analise_prioridade = context.get('analise_prioridade', [])
    if analise_prioridade:
        adicionar_titulo_secao(doc, "Análise por Prioridade")
        _montar_tabela_analise(doc, "Prioridade", analise_prioridade)
        doc.add_paragraph()

    # ===== DETALHAMENTO =====
    adicionar_titulo_secao(doc, "Detalhamento das Tarefas")
    _montar_tabela_detalhamento(doc, context.get('tarefas', []))

    # ===== RESPONSE =====
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = (
        f'attachment; filename="relatorio_tarefas_{timezone.now():%Y%m%d_%H%M}.docx"'
    )
    return response


# ============================================================
# HELPERS INTERNOS DOCX
# ============================================================
def _montar_tabela_analise(doc, label_coluna, itens):
    """Monta tabela de análise (Status ou Prioridade)."""
    tabela = doc.add_table(rows=1, cols=3)

    # Cabeçalho
    hdr = tabela.rows[0].cells
    hdr[0].text = label_coluna
    hdr[1].text = "Quantidade"
    hdr[2].text = "Percentual"

    # Linhas
    for item in itens:
        row = tabela.add_row().cells
        row[0].text = item['label']
        row[1].text = str(item['quantidade'])
        row[2].text = f"{item['percentual']}%"

    # 1️⃣ Aplica estilo base (zebra/bordas) PRIMEIRO
    aplicar_estilo_tabela_word(tabela, larguras_cm=[6, 4, 4])

    # 2️⃣ Aplica badges DEPOIS (preserva cores)
    for idx, item in enumerate(itens, start=1):
        adicionar_badge_cell(
            tabela.rows[idx].cells[0],
            item['label'],
            item.get('slug', '')
        )


def _montar_tabela_detalhamento(doc, tarefas):
    """Monta tabela de detalhamento das tarefas."""
    headers = ["#", "Título", "Responsável", "Status", "Prioridade",
               "Projeto", "Criação", "Prazo"]
    tabela = doc.add_table(rows=1, cols=len(headers))

    # Cabeçalho
    for idx, h in enumerate(headers):
        tabela.rows[0].cells[idx].text = h

    if not tarefas:
        # Linha "vazio"
        row = tabela.add_row().cells
        merged = row[0].merge(row[-1])
        merged.text = "Nenhuma tarefa encontrada."
        aplicar_estilo_tabela_word(
            tabela,
            larguras_cm=[1, 4.5, 3, 2.5, 2.5, 3, 2, 2]
        )
        return

    # Preenche linhas (sem badges ainda)
    for i, t in enumerate(tarefas, start=1):
        row = tabela.add_row().cells
        row[0].text = str(i)
        row[1].text = t.titulo or "—"
        row[2].text = (
            (t.responsavel.get_full_name() or t.responsavel.username)
            if t.responsavel else "—"
        )
        row[3].text = t.get_status_display()
        row[4].text = t.get_prioridade_display()
        row[5].text = str(t.projeto) if t.projeto else "—"
        row[6].text = t.data_criacao.strftime("%d/%m/%Y") if t.data_criacao else "—"
        row[7].text = t.prazo.strftime("%d/%m/%Y") if t.prazo else "—"

    # 1️⃣ Estilo base
    aplicar_estilo_tabela_word(
        tabela,
        larguras_cm=[1, 4.5, 3, 2.5, 2.5, 3, 2, 2]
    )

    # 2️⃣ Badges depois (preserva cores)
    for i, t in enumerate(tarefas, start=1):
        cells = tabela.rows[i].cells
        adicionar_badge_cell(cells[3], t.get_status_display(), t.status)
        adicionar_badge_cell(cells[4], t.get_prioridade_display(), t.prioridade)
