"""
Estilos padronizados para exportação Word — mesma identidade visual.
"""
from docx.shared import Pt, RGBColor, Cm, Mm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ===== CORES (mesmas do PDF/Excel) =====
COR_PRIMARIA = RGBColor(0x0D, 0x6E, 0xFD)
COR_TEXTO = RGBColor(0x21, 0x25, 0x29)
COR_MUTED = RGBColor(0x6C, 0x75, 0x7D)
COR_BRANCO = RGBColor(0xFF, 0xFF, 0xFF)

HEX_PRIMARIA = "0D6EFD"
HEX_ZEBRA    = "F8F9FA"
HEX_BORDA    = "DEE2E6"

# Status / Prioridade (fundo dos badges)
HEX_BADGE = {
    "pendente":   "FFF3CD",
    "andamento":  "CFE2FF",
    "concluida":  "D1E7DD",
    "concluido":  "D1E7DD",
    "atrasada":   "F8D7DA",
    "cancelada":  "E2E3E5",
    "baixa":      "FFF3CD",
    "media":      "CFE2FF",
    "alta":       "F8D7DA",
    "urgente":    "E2E3E5",
}

# KPIs
KPI_CORES = {
    "info":    "0D6EFD",
    "success": "198754",
    "warning": "FFC107",
    "danger":  "DC3545",
}


# ============================================================
# CABEÇALHO
# ============================================================
def aplicar_cabecalho_word(doc, titulo, subtitulo, data_emissao):
    """Cabeçalho institucional padrão."""
    # Título
    p = doc.add_paragraph()
    run = p.add_run(titulo)
    run.font.size = Pt(18)
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.color.rgb = COR_PRIMARIA

    # Subtítulo
    p = doc.add_paragraph()
    run = p.add_run(subtitulo)
    run.font.size = Pt(10)
    run.font.name = "Calibri"
    run.font.color.rgb = COR_MUTED

    # Data
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run(f"Emitido em: {data_emissao}")
    run.font.size = Pt(9)
    run.font.name = "Calibri"
    run.font.color.rgb = COR_MUTED

    # Linha separadora azul
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '18')
    bottom.set(qn('w:color'), HEX_PRIMARIA)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================================================
# TÍTULO DE SEÇÃO
# ============================================================
def adicionar_titulo_secao(doc, texto):
    """Título de seção com underline azul."""
    p = doc.add_paragraph()
    run = p.add_run(texto)
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.name = "Calibri"
    run.font.color.rgb = COR_PRIMARIA

    # Borda inferior
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:color'), HEX_BORDA)
    pBdr.append(bottom)
    pPr.append(pBdr)


# ============================================================
# KPIs (4 cards lado a lado)
# ============================================================
def adicionar_kpis(doc, kpis):
    """
    Adiciona linha de KPIs.
    kpis = [
        {"label": "Total", "value": 42, "tipo": "info"},
        {"label": "Concluídas", "value": 18, "tipo": "success"},
        ...
    ]
    """
    table = doc.add_table(rows=1, cols=len(kpis))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for idx, kpi in enumerate(kpis):
        cell = table.rows[0].cells[idx]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Fundo cinza claro
        _set_cell_shading(cell, HEX_ZEBRA)
        # Borda esquerda colorida (simulada com borda mais grossa)
        _set_cell_left_border(cell, KPI_CORES.get(kpi.get("tipo", "info"), HEX_PRIMARIA))

        # Limpa parágrafo padrão
        cell.paragraphs[0].text = ""

        # Label
        p = cell.paragraphs[0]
        run = p.add_run(kpi["label"].upper())
        run.font.size = Pt(8)
        run.font.bold = True
        run.font.name = "Calibri"
        run.font.color.rgb = COR_MUTED

        # Valor
        p2 = cell.add_paragraph()
        run = p2.add_run(str(kpi["value"]))
        run.font.size = Pt(20)
        run.font.bold = True
        run.font.name = "Calibri"
        run.font.color.rgb = COR_TEXTO


# ============================================================
# TABELA PADRÃO
# ============================================================
def aplicar_estilo_tabela_word(table, larguras_cm=None):
    """Aplica estilo padrão na tabela (cabeçalho azul + zebra)."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Cabeçalho azul
    for cell in table.rows[0].cells:
        _set_cell_shading(cell, HEX_PRIMARIA)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        for para in cell.paragraphs:
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = COR_BRANCO
                run.font.size = Pt(10)
                run.font.name = "Calibri"

    # Zebra
    for i, row in enumerate(table.rows[1:], start=1):
        if i % 2 == 0:
            for cell in row.cells:
                _set_cell_shading(cell, HEX_ZEBRA)
        for cell in row.cells:
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.font.size is None:
                        run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    # Só seta cor se ainda não tiver
                    try:
                        if run.font.color.rgb is None:
                            run.font.color.rgb = COR_TEXTO
                    except AttributeError:
                        run.font.color.rgb = COR_TEXTO


    # Larguras
    if larguras_cm:
        for row in table.rows:
            for idx, largura in enumerate(larguras_cm):
                if idx < len(row.cells):
                    row.cells[idx].width = Cm(largura)

    # Bordas em todas as células
    _aplicar_bordas_tabela(table)


def adicionar_badge_cell(cell, texto, slug):
    """Adiciona texto estilizado como badge dentro de uma célula."""
    cell.paragraphs[0].text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(texto)
    run.font.size = Pt(8)
    run.font.bold = True
    run.font.name = "Calibri"
    cor_fundo = HEX_BADGE.get(slug, "E2E3E5")
    _set_cell_shading(cell, cor_fundo)


# ============================================================
# HELPERS XML INTERNOS
# ============================================================
def _set_cell_shading(cell, hex_color):
    """Define cor de fundo da célula."""
    tcPr = cell._tc.get_or_add_tcPr()
    # Remove shading anterior se existir
    for shd in tcPr.findall(qn('w:shd')):
        tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _set_cell_left_border(cell, hex_color, sz="24"):
    """Define borda esquerda colorida (para KPI cards)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), sz)
    left.set(qn('w:color'), hex_color)
    tcBorders.append(left)


def _aplicar_bordas_tabela(table):
    """Aplica bordas finas em toda a tabela."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)

    tblBorders = OxmlElement('w:tblBorders')
    for borda in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        b = OxmlElement(f'w:{borda}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), HEX_BORDA)
        tblBorders.append(b)

    # Remove bordas antigas
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblPr.append(tblBorders)
