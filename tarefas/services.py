# tarefas/services.py

"""
Serviço de relatórios do módulo de tarefas.
Funções para preparar contexto e exportar em PDF, CSV e DOCX.
"""

import csv
import io
import logging

from django.db.models import Count
from django.http import HttpResponse
from django.template.loader import render_to_string
from django.utils import timezone

logger = logging.getLogger(__name__)


# =============================================================================
# CONTEXTO DO RELATÓRIO
# =============================================================================

def preparar_contexto_relatorio(queryset):
    """
    Prepara dados estatísticos a partir do queryset de tarefas
    para uso no template de relatório, dashboard e exportações.
    """
    from .models import Tarefas

    total = queryset.count()

    # Dicionários de labels legíveis
    status_labels = dict(Tarefas.STATUS_CHOICES)
    prioridade_labels = dict(Tarefas.PRIORIDADE_CHOICES)

    # ═══ Contagem por STATUS ═══
    status_counts = (
        queryset
        .values('status')
        .annotate(total=Count('id'))
        .order_by('status')
    )
    status_data = []
    for item in status_counts:
        key = item['status']
        count = item['total']
        pct = round((count / total * 100), 1) if total > 0 else 0
        status_data.append({
            'key': key,                                # chave técnica (para CSS/JS)
            'status': key,                             # ← compatibilidade com gráficos
            'label': status_labels.get(key, key),      # nome legível
            'total': count,
            'percentual': pct,
        })

    # ═══ Contagem por PRIORIDADE ═══
    prioridade_counts = (
        queryset
        .values('prioridade')
        .annotate(total=Count('id'))
        .order_by('prioridade')
    )
    prioridade_data = []
    for item in prioridade_counts:
        key = item['prioridade']
        count = item['total']
        pct = round((count / total * 100), 1) if total > 0 else 0
        prioridade_data.append({
            'key': key,                                    # chave técnica
            'prioridade': key,                             # ← compatibilidade com gráficos
            'label': prioridade_labels.get(key, key),      # nome legível
            'total': count,
            'percentual': pct,
        })

    # Tarefas atrasadas
    agora = timezone.now()
    atrasadas = queryset.filter(prazo__lt=agora).exclude(
        status__in=['concluida', 'cancelada']
    ).count()

    # Concluídas
    concluidas = queryset.filter(status='concluida').count()

    return {
        'total_tarefas': total,
        'status_data': status_data,
        'prioridade_data': prioridade_data,
        'atrasadas': atrasadas,
        'concluidas': concluidas,
        'tarefas': queryset,
    }


# =============================================================================
# EXPORTAÇÃO PDF
# =============================================================================

def gerar_pdf_relatorio(context):
    """
    Gera um relatório em PDF usando WeasyPrint (prioridade) ou xhtml2pdf (fallback).
    """
    try:
        from weasyprint import HTML

        html_string = render_to_string('tarefas/relatorios/relatorio_pdf.html', context)
        pdf_file = HTML(string=html_string).write_pdf()

    except ImportError:
        try:
            from xhtml2pdf import pisa

            html_string = render_to_string('tarefas/relatorios/relatorio_pdf.html', context)
            result = io.BytesIO()
            pisa_status = pisa.CreatePDF(io.StringIO(html_string), dest=result)
            if pisa_status.err:
                logger.error("Erro ao gerar PDF com xhtml2pdf.")
                return HttpResponse('Erro ao gerar PDF.', status=500)
            pdf_file = result.getvalue()

        except ImportError:
            logger.error("Nenhuma biblioteca de PDF instalada (weasyprint ou xhtml2pdf).")
            return HttpResponse(
                'Nenhuma biblioteca de PDF instalada (weasyprint ou xhtml2pdf).',
                status=500,
            )

    response = HttpResponse(pdf_file, content_type='application/pdf')
    response['Content-Disposition'] = 'attachment; filename="relatorio_tarefas.pdf"'
    return response


# =============================================================================
# EXPORTAÇÃO CSV
# =============================================================================

def gerar_csv_relatorio(context):
    """
    Gera um relatório em CSV a partir do contexto.
    """
    response = HttpResponse(content_type='text/csv; charset=utf-8')
    response['Content-Disposition'] = 'attachment; filename="relatorio_tarefas.csv"'
    response.write('\ufeff')  # BOM para Excel reconhecer UTF-8

    writer = csv.writer(response, delimiter=';')
    writer.writerow([
        'ID', 'Título', 'Status', 'Prioridade',
        'Responsável', 'Data Criação', 'Prazo',
    ])

    for tarefa in context.get('tarefas', []):
        writer.writerow([
            tarefa.id,
            tarefa.titulo,
            tarefa.get_status_display() if hasattr(tarefa, 'get_status_display') else tarefa.status,
            tarefa.get_prioridade_display() if hasattr(tarefa, 'get_prioridade_display') else tarefa.prioridade,
            str(tarefa.responsavel) if tarefa.responsavel else '-',
            tarefa.data_criacao.strftime('%d/%m/%Y %H:%M') if tarefa.data_criacao else '-',
            tarefa.prazo.strftime('%d/%m/%Y %H:%M') if tarefa.prazo else '-',
        ])

    return response


# =============================================================================
# EXPORTAÇÃO DOCX
# =============================================================================

def gerar_docx_relatorio(context):
    """
    Gera um relatório em DOCX usando python-docx.
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error("Biblioteca python-docx não instalada.")
        return HttpResponse(
            'Biblioteca python-docx não instalada.',
            status=500,
        )

    doc = Document()
    doc.add_heading('Relatório de Tarefas', level=1)

    # Resumo
    total = context.get('total_tarefas', 0)
    concluidas = context.get('concluidas', 0)
    atrasadas = context.get('atrasadas', 0)
    doc.add_paragraph(
        f"Total: {total}  ·  Concluídas: {concluidas}  ·  Atrasadas: {atrasadas}"
    )

    # ═══ Análise por Status ═══
    status_data = context.get('status_data', [])
    if status_data:
        doc.add_heading('Análise por Status', level=2)
        for item in status_data:
            doc.add_paragraph(
                f"{item['label']}: {item['total']} ({item['percentual']}%)"
            )

    # ═══ Análise por Prioridade ═══
    prioridade_data = context.get('prioridade_data', [])
    if prioridade_data:
        doc.add_heading('Análise por Prioridade', level=2)
        for item in prioridade_data:
            doc.add_paragraph(
                f"{item['label']}: {item['total']} ({item['percentual']}%)"
            )

    doc.add_paragraph('')

    # Tabela de tarefas
    tarefas = context.get('tarefas', [])
    if tarefas:
        doc.add_heading('Detalhamento', level=2)
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Light Grid Accent 1'

        header_cells = table.rows[0].cells
        headers = ['Título', 'Status', 'Prioridade', 'Responsável', 'Prazo']
        for i, header in enumerate(headers):
            header_cells[i].text = header

        for tarefa in tarefas:
            row_cells = table.add_row().cells
            row_cells[0].text = tarefa.titulo or ''
            row_cells[1].text = (
                tarefa.get_status_display()
                if hasattr(tarefa, 'get_status_display')
                else tarefa.status or ''
            )
            row_cells[2].text = (
                tarefa.get_prioridade_display()
                if hasattr(tarefa, 'get_prioridade_display')
                else tarefa.prioridade or ''
            )
            row_cells[3].text = str(tarefa.responsavel) if tarefa.responsavel else '-'
            row_cells[4].text = tarefa.prazo.strftime('%d/%m/%Y %H:%M') if tarefa.prazo else '-'

    # Gera response
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
    )
    response['Content-Disposition'] = 'attachment; filename="relatorio_tarefas.docx"'
    return response


# ═══════════════════════════════════════════════════════════════════════
# Re-exportar funções de histórico para manter imports simples
# ═══════════════════════════════════════════════════════════════════════

from .historico_services import (  # noqa: F401, E402
    registrar_criacao_tarefa,
    registrar_alteracoes_tarefa,
    registrar_alteracao_status,
    registrar_alteracao_participantes,
)

