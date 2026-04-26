
# automovel/views.py

import io
import json
from datetime import timedelta, datetime
from core.mixins import AppPermissionMixin, ViewFilialScopedMixin, FuncionarioRequiredMixin
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth.mixins import PermissionRequiredMixin
from django.contrib.messages.views import SuccessMessageMixin
from django.core.exceptions import ObjectDoesNotExist, PermissionDenied
from django.db.models import Q
from django.http import (
    Http404, HttpResponse, HttpResponseBadRequest, JsonResponse,
)
from django.shortcuts import get_object_or_404, redirect, render
from django.urls import NoReverseMatch, reverse, reverse_lazy
from django.utils import timezone
from django.utils.decorators import method_decorator
from django.views.decorators.csrf import csrf_exempt
from django.views.generic import (
    CreateView, DeleteView, DetailView, ListView, TemplateView, UpdateView, View,
)

# Bibliotecas de terceiros
from docx import Document, settings
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Inches
from openpyxl import Workbook
from openpyxl.cell import MergedCell
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter
from departamento_pessoal.models import Funcionario
from .models import Filial
from .forms import AgendamentoForm, CarroForm, ChecklistForm, ManutencaoForm
from .models import (
    Carro, Carro_agendamento, Carro_checklist,
    Carro_manutencao, Carro_rastreamento,
)


# ═══════════════════════════════════════════════════════════════════════════════
# MIXIN CENTRAL DE FILIAL
# ═══════════════════════════════════════════════════════════════════════════════

class FilialAtivaMixin:
    """Fornece utilitários para obter/filtrar pela filial ativa do usuário."""

    def get_filial_ativa(self):
        filial_id = self.request.session.get('active_filial_id')
        if filial_id:
            try:
                return Filial.objects.get(pk=filial_id)
            except Filial.DoesNotExist:
                pass

        user = self.request.user
        filial_ativa = getattr(user, 'filial_ativa', None)
        if filial_ativa:
            return filial_ativa

        try:
            funcionario = Funcionario.objects.select_related('filial').get(usuario=user)
            return funcionario.filial
        except Funcionario.DoesNotExist:
            pass

        return None

    def get_filial_ativa_id(self):
        filial = self.get_filial_ativa()
        return filial.id if filial else None

    def filter_queryset_by_filial(self, queryset):
        filial = self.get_filial_ativa()
        if filial:
            return queryset.filter(filial=filial)
        elif not self.request.user.is_superuser:
            return queryset.none()
        return queryset


# ═══════════════════════════════════════════════════════════════════════════════
# MIXIN DE VISIBILIDADE POR PERFIL
# ═══════════════════════════════════════════════════════════════════════════════

class AutomovelVisibilityMixin(FilialAtivaMixin):
    """
    Controla visibilidade dos registros do app Automóvel.

    Regras:
    - Superuser           → tudo
    - view_all_automovel  → tudo da filial ativa
    - Usuário comum       → apenas registros onde é owner (owner_field)

    IMPORTANTE: se a view usa um model que NÃO tem campo 'usuario' (ex.: Carro),
    defina `apply_owner_filter = False` para não quebrar.
    """

    owner_field = 'usuario'
    apply_owner_filter = True  # False para models sem campo de "dono"

    def apply_visibility(self, queryset):
        user = self.request.user

        if user.is_superuser:
            return queryset

        if user.has_perm('automovel.view_all_automovel'):
            return queryset

        if not self.apply_owner_filter:
            # Model sem "dono" (ex.: Carro) → sem filtro adicional (só filial scope)
            return queryset

        return queryset.filter(**{self.owner_field: user})


# ═══════════════════════════════════════════════════════════════════════════════
# MIXIN BASE
# ═══════════════════════════════════════════════════════════════════════════════

class AutomovelBaseMixin(FuncionarioRequiredMixin, AppPermissionMixin, AutomovelVisibilityMixin, ViewFilialScopedMixin,):
    """
    Mixin base para views do módulo Automóvel.

    Regras de acesso:
      1. Anônimo → LoginRequiredMixin redireciona pro login
      2. Superuser → bypass total
      3. Usuário sem Funcionario → tela amigável (core:sem_funcionario)
      4. Usuário com Funcionario → valida permissões específicas
    """

    login_url = 'login'
    app_label_required = 'automovel'
    modulo_nome = 'Automóvel' 

    def dispatch(self, request, *args, **kwargs):
        # FuncionarioRequiredMixin já trata: anônimo, superuser, sem funcionario
        # Aqui só validamos a permissão específica do módulo
        response = super().dispatch(request, *args, **kwargs)

        # Se o super já retornou redirect (sem funcionario, login etc.), respeita
        if response.status_code in (301, 302):
            return response

        return response

    def has_permission(self):
        """Hook chamado pelo AppPermissionMixin (se for o caso)."""
        funcionario = getattr(self.request, 'funcionario', None)
        if self.request.user.is_superuser:
            return True
        return self.tem_permissao_automovel(funcionario)

    def tem_permissao_automovel(self, funcionario):
        if funcionario is None:
            return self.request.user.is_superuser

        setor = getattr(funcionario, 'setor', None)
        if setor and setor.nome in ['TI', 'Frota', 'Diretoria']:
            return True

        cargo = getattr(funcionario, 'cargo', None)
        if cargo and getattr(cargo, 'permite_automovel', False):
            return True

        usuario = getattr(funcionario, 'usuario', None)
        if usuario:
            perms_automovel = [
                p for p in usuario.get_all_permissions()
                if p.startswith('automovel.')
            ]
            if perms_automovel:
                return True

        return False

# ═══════════════════════════════════════════════════════════════════════════════
# MIXINS LOCAIS DO APP (reutilizáveis)
# ═══════════════════════════════════════════════════════════════════════════════

class ChecklistFormSectionsMixin:
    """Organiza os campos do checklist em seções para o template."""

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        form = context['form']
        context['form_sections'] = [
            {
                'id': 'frontal',
                'title': 'Vistoria da Parte Frontal',
                'status_field': form['revisao_frontal_status'],
                'photo_field': form['foto_frontal'],
            },
            {
                'id': 'traseira',
                'title': 'Vistoria da Parte Traseira',
                'status_field': form['revisao_trazeira_status'],
                'photo_field': form['foto_trazeira'],
            },
            {
                'id': 'motorista',
                'title': 'Vistoria do Lado do Motorista',
                'status_field': form['revisao_lado_motorista_status'],
                'photo_field': form['foto_lado_motorista'],
            },
            {
                'id': 'passageiro',
                'title': 'Vistoria do Lado do Passageiro',
                'status_field': form['revisao_lado_passageiro_status'],
                'photo_field': form['foto_lado_passageiro'],
            },
        ]
        return context


class _ReverseGeocodeMixin:
    """Geocoding reverso via OpenStreetMap Nominatim."""

    def _reverse_geocode(self, lat, lng):
        try:
            import requests as http_requests
            headers = {'User-Agent': 'ControleTarefas/1.0 (esg@cetestsp.com.br)'}
            response = http_requests.get(
                f'https://nominatim.openstreetmap.org/reverse'
                f'?format=json&lat={lat}&lon={lng}&zoom=18',
                headers=headers, timeout=5,
            )
            if response.status_code == 200:
                return response.json().get('display_name', 'Endereço não identificado')
        except Exception:
            pass
        return 'Endereço não disponível'


# ═══════════════════════════════════════════════════════════════════════════════
# DASHBOARD E CALENDÁRIO
# ═══════════════════════════════════════════════════════════════════════════════

class DashboardView(AutomovelBaseMixin, ListView):
    template_name = 'automovel/dashboard.html'
    context_object_name = 'ultimos_agendamentos'
    model = Carro_agendamento

    def get_queryset(self):
        qs = (
            Carro_agendamento.objects.for_request(self.request)
            .select_related('carro', 'usuario')
        )
        qs = self.apply_visibility(qs)
        return qs.order_by('-data_hora_agenda')[:5]

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        hoje = timezone.now().date()

        # Carro não tem campo 'usuario' → desabilita filtro de owner
        carros_qs = Carro.objects.for_request(self.request).filter(ativo=True)

        agendamentos_qs = self.apply_visibility(
            Carro_agendamento.objects.for_request(self.request)
        )

        context['total_carros'] = carros_qs.count()
        context['carros_disponiveis'] = carros_qs.filter(disponivel=True).count()
        context['agendamentos_hoje'] = agendamentos_qs.filter(
            data_hora_agenda__date=hoje, status='agendado'
        ).count()
        context['manutencao_proxima'] = carros_qs.filter(
            data_proxima_manutencao__lte=hoje + timedelta(days=7),
            data_proxima_manutencao__gte=hoje,
        ).count()
        context['filial_ativa'] = self.get_filial_ativa()
        return context


class CalendarioView(AutomovelBaseMixin, TemplateView):
    template_name = 'automovel/calendario.html'

    def get_context_data(self, **kwargs):
        ctx = super().get_context_data(**kwargs)
        ctx['filial_ativa'] = self.get_filial_ativa()
        return ctx


class CalendarioAPIView(AutomovelBaseMixin, View):
    STATUS_COLORS = {
        'agendado': '#0d6efd',
        'em_andamento': '#ffc107',
        'finalizado': '#198754',
        'atrasado': '#dc3545',
    }

    def get(self, request, *args, **kwargs):
        eventos = self._get_eventos_agendamentos(request)
        eventos += self._get_eventos_manutencoes(request)
        return JsonResponse(eventos, safe=False)

    def _get_eventos_agendamentos(self, request):
        qs = (
            Carro_agendamento.objects.for_request(request)
            .filter(cancelar_agenda=False)
            .select_related('carro')
        )
        qs = self.apply_visibility(qs)

        eventos = []
        for ag in qs:
            start = ag.data_hora_agenda
            end = ag.data_hora_devolucao or start + timedelta(hours=1)
            eventos.append({
                'id': f"ag_{ag.id}",
                'title': f"{ag.carro.placa} - {ag.funcionario or 'Sem Motorista'}",
                'start': start.isoformat(),
                'end': end.isoformat(),
                'url': reverse('automovel:agendamento_detail', kwargs={'pk': ag.id}),
                'color': self.STATUS_COLORS.get(ag.status, '#6c757d'),
                'extendedProps': {'tipo': 'agendamento', 'status': ag.get_status_display()},
            })
        return eventos

    def _get_eventos_manutencoes(self, request):
        qs = (
            Carro_manutencao.objects.for_request(request)
            .select_related('carro')
        )
        qs = self.apply_visibility(qs)

        eventos = []
        for man in qs:
            cor = '#6c757d' if man.concluida else '#d63384'
            prefixo = '✅ (OK) ' if man.concluida else ''
            eventos.append({
                'id': f"man_{man.id}",
                'title': f"{prefixo}🔧 {man.get_tipo_display()} - {man.carro.placa}",
                'start': man.data_manutencao.isoformat(),
                'allDay': True,
                'color': cor,
                'url': '#',
                'extendedProps': {'tipo': 'manutencao', 'descricao': man.descricao},
            })
        return eventos


# ═══════════════════════════════════════════════════════════════════════════════
# CARRO CRUD
# ═══════════════════════════════════════════════════════════════════════════════

class CarroListView(AutomovelBaseMixin, PermissionRequiredMixin, ListView):
    """Carro não tem dono → apply_owner_filter = False"""
    model = Carro
    template_name = 'automovel/carro_list.html'
    context_object_name = 'carros'
    paginate_by = 10
    permission_required = 'automovel.view_carro'
    apply_owner_filter = False

    def get_queryset(self):
        qs = Carro.objects.for_request(self.request).filter(ativo=True)
        qs = self.apply_visibility(qs)
        search = self.request.GET.get('search')
        if search:
            qs = qs.filter(
                Q(placa__icontains=search)
                | Q(modelo__icontains=search)
                | Q(marca__icontains=search)
            )
        return qs


class CarroCreateView(AutomovelBaseMixin, PermissionRequiredMixin, CreateView):
    model = Carro
    form_class = CarroForm
    template_name = 'automovel/carro_form.html'
    success_url = reverse_lazy('automovel:carro_list')
    success_message = "Carro cadastrado com sucesso!"
    apply_owner_filter = False
    permission_required = 'automovel.add_carro'

    def form_valid(self, form):
        filial = self.get_filial_ativa()
        if not filial:
            messages.error(self.request, "Nenhuma filial ativa. Selecione uma filial no menu superior.")
            return self.form_invalid(form)
        form.instance.filial = filial
        return super().form_valid(form)


class CarroUpdateView(AutomovelBaseMixin, SuccessMessageMixin, UpdateView):
    model = Carro
    form_class = CarroForm
    template_name = 'automovel/carro_form.html'
    success_url = reverse_lazy('automovel:carro_list')
    success_message = "Carro atualizado com sucesso!"
    apply_owner_filter = False


class CarroDetailView(AutomovelBaseMixin, DetailView):
    model = Carro
    template_name = 'automovel/carro_detail.html'
    context_object_name = "carro"
    apply_owner_filter = False

    def get_queryset(self):
        return super().get_queryset().select_related('filial').prefetch_related(
            'agendamentos', 'manutencoes'
        )

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        carro = self.object

        # Filtra agendamentos por visibilidade (aqui SIM tem 'usuario')
        agendamentos_qs = carro.agendamentos.all().select_related('usuario')
        # Reativa o filtro de owner temporariamente
        original = self.apply_owner_filter
        self.apply_owner_filter = True
        agendamentos_qs = self.apply_visibility(agendamentos_qs)
        self.apply_owner_filter = original

        context['agendamentos'] = agendamentos_qs.order_by('-data_hora_agenda')[:10]
        context['manutencoes'] = carro.manutencoes.all().order_by('-data_manutencao')[:10]
        context['filial_ativa'] = self.get_filial_ativa()
        return context

class CarroDeleteView(AutomovelBaseMixin, PermissionRequiredMixin, DeleteView):
    permission_required = 'automovel.delete_carro'

    def post(self, request, pk):
        carro = get_object_or_404(Carro.objects.for_request(request), pk=pk)
        carro.ativo = False
        carro.save(update_fields=['ativo'])
        messages.success(request, "Carro desativado com sucesso!")
        return redirect('automovel:carro_list') 
    

# ═══════════════════════════════════════════════════════════════════════════════
# AGENDAMENTO CRUD
# ═══════════════════════════════════════════════════════════════════════════════

class AgendamentoListView(AutomovelBaseMixin, ListView):
    model = Carro_agendamento
    template_name = 'automovel/agendamento_list.html'
    context_object_name = 'agendamentos'
    paginate_by = 20

    def get_queryset(self):
        qs = super().get_queryset().select_related('carro', 'usuario')
        return self.apply_visibility(qs)


class AgendamentoCreateView(AutomovelBaseMixin, SuccessMessageMixin, CreateView):
    model = Carro_agendamento
    form_class = AgendamentoForm
    template_name = 'automovel/agendamento_form.html'
    success_url = reverse_lazy('automovel:agendamento_list')
    success_message = "Agendamento criado com sucesso!"

    def form_valid(self, form):
        filial = self.get_filial_ativa()
        if not filial:
            messages.error(self.request, "Nenhuma filial ativa selecionada.")
            return self.form_invalid(form)
        form.instance.usuario = self.request.user
        form.instance.filial = filial
        return super().form_valid(form)


class AgendamentoUpdateView(AutomovelBaseMixin, SuccessMessageMixin, UpdateView):
    model = Carro_agendamento
    form_class = AgendamentoForm
    template_name = 'automovel/agendamento_form.html'
    success_message = "Agendamento atualizado com sucesso!"

    def get_queryset(self):
        return self.apply_visibility(super().get_queryset())

    def get_success_url(self):
        return reverse('automovel:agendamento_detail', kwargs={'pk': self.object.pk})


class AgendamentoDetailView(AutomovelBaseMixin, DetailView):
    model = Carro_agendamento
    template_name = 'automovel/agendamento_detail.html'

    def get_queryset(self):
        return self.apply_visibility(super().get_queryset())

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        agendamento = self.object

        historico_qs = (
            Carro_agendamento.objects.for_request(self.request)
            .filter(carro=agendamento.carro)
            .exclude(id=agendamento.id)
            .select_related('usuario')
        )
        historico_qs = self.apply_visibility(historico_qs)
        context['historico_agendamentos'] = historico_qs.order_by('-data_hora_agenda')[:10]

        context['rastreamentos'] = (
            Carro_rastreamento.objects.filter(agendamento=agendamento)
            .order_by('-data_hora')[:50]
        )
        context['checklist_saida'] = agendamento.checklists.filter(tipo='saida').first()
        context['checklist_retorno'] = agendamento.checklists.filter(tipo='retorno').first()
        context['mapbox_access_token'] = getattr(settings, 'MAPBOX_ACCESS_TOKEN', '')
        context['filial_ativa'] = self.get_filial_ativa()
        return context


class AgendamentoFinalizarView(AutomovelBaseMixin, View):

    def post(self, request, pk):
        base_qs = Carro_agendamento.objects.for_request(request)
        base_qs = self.apply_visibility(base_qs)
        agendamento = get_object_or_404(base_qs, pk=pk)

        km_final = request.POST.get('km_final')
        observacoes = request.POST.get('observacoes_devolucao')

        if not km_final:
            messages.error(request, "A Quilometragem Final é obrigatória.")
            return redirect('automovel:carro_detail', pk=agendamento.carro.pk)

        try:
            km_final_float = float(km_final)
        except ValueError:
            messages.error(request, "Valor de KM inválido.")
            return redirect('automovel:carro_detail', pk=agendamento.carro.pk)

        if km_final_float < agendamento.carro.quilometragem:
            messages.error(
                request,
                f"Erro: O KM informado ({km_final_float}) é menor que o KM "
                f"atual do veículo ({agendamento.carro.quilometragem})."
            )
            return redirect('automovel:carro_detail', pk=agendamento.carro.pk)

        agendamento.km_final = km_final_float
        agendamento.ocorrencia = observacoes
        agendamento.status = 'finalizado'
        agendamento.data_hora_devolucao = timezone.now()
        agendamento.save()

        carro = agendamento.carro
        carro.quilometragem = km_final_float
        carro.disponivel = True
        carro.save(update_fields=['quilometragem', 'disponivel'])

        messages.success(request, "Agendamento finalizado e veículo devolvido com sucesso!")
        return redirect('automovel:carro_detail', pk=carro.pk)


# ═══════════════════════════════════════════════════════════════════════════════
# CHECKLIST CRUD
# ═══════════════════════════════════════════════════════════════════════════════

class ChecklistListView(AutomovelBaseMixin, ListView):
    model = Carro_checklist
    template_name = 'automovel/checklist_list.html'
    context_object_name = 'checklists'

    def get_queryset(self):
        qs = super().get_queryset().select_related('agendamento__carro', 'usuario')
        return self.apply_visibility(qs)


class ChecklistCreateView(
    AutomovelBaseMixin, SuccessMessageMixin,
    ChecklistFormSectionsMixin, CreateView
):
    model = Carro_checklist
    form_class = ChecklistForm
    template_name = 'automovel/checklist_form.html'

    def setup(self, request, *args, **kwargs):
        super().setup(request, *args, **kwargs)
        self.agendamento = None
        self.tipo_checklist = request.GET.get('tipo', 'saida')

    def dispatch(self, request, *args, **kwargs):
        # Dispara o guard do AutomovelBaseMixin primeiro
        parent = super().dispatch(request, *args, **kwargs)
        # Se já retornou 403 do guard, não continua
        if hasattr(parent, 'status_code') and parent.status_code == 403:
            return parent
        return parent

    def get_agendamento(self):
        """Busca o agendamento validando visibilidade."""
        if self.agendamento is None:
            agendamento_qs = (
                Carro_agendamento.objects.for_request(self.request)
                .select_related('carro')
            )
            agendamento_qs = self.apply_visibility(agendamento_qs)
            self.agendamento = get_object_or_404(
                agendamento_qs, pk=self.kwargs.get('agendamento_pk')
            )
        return self.agendamento

    def get(self, request, *args, **kwargs):
        agendamento = self.get_agendamento()
        if Carro_checklist.objects.filter(
            agendamento=agendamento, tipo=self.tipo_checklist
        ).exists():
            messages.error(
                request,
                f"Um checklist de '{self.tipo_checklist}' já existe para este agendamento."
            )
            return redirect('automovel:agendamento_detail', pk=agendamento.pk)
        return super().get(request, *args, **kwargs)

    def post(self, request, *args, **kwargs):
        self.get_agendamento()
        return super().post(request, *args, **kwargs)

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['tipo_checklist'] = self.tipo_checklist
        return kwargs

    def get_initial(self):
        return {'agendamento': self.get_agendamento(), 'tipo': self.tipo_checklist}

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['agendamento'] = self.get_agendamento()
        return context

    def form_valid(self, form):
        usuario = self.request.user
        try:
            funcionario_logado = usuario.funcionario
        except ObjectDoesNotExist:
            messages.error(
                self.request,
                "Seu usuário não está vinculado a um Funcionário. Contate o RH."
            )
            return self.form_invalid(form)

        agendamento = self.get_agendamento()

        if self.tipo_checklist == 'saida':
            km_inicial = form.cleaned_data.get('km_inicial')
            if not km_inicial:
                form.add_error('km_inicial', 'A quilometragem inicial é obrigatória.')
                return self.form_invalid(form)
            agendamento.km_inicial = km_inicial
            agendamento.save(update_fields=['km_inicial'])

        elif self.tipo_checklist == 'retorno':
            km_final = form.cleaned_data.get('km_final')
            if not km_final:
                form.add_error('km_final', 'A quilometragem final é obrigatória.')
                return self.form_invalid(form)
            agendamento.km_final = km_final
            agendamento.save(update_fields=['km_final'])

        form.instance.agendamento = agendamento
        form.instance.usuario = usuario
        form.instance.responsavel = funcionario_logado
        form.instance.filial = funcionario_logado.filial

        self.success_message = f"Checklist de {self.tipo_checklist} registrado com sucesso!"
        if self.tipo_checklist == 'retorno':
            self.success_message += " Agendamento finalizado!"

        return super().form_valid(form)

    def get_success_url(self):
        return reverse(
            'automovel:agendamento_detail',
            kwargs={'pk': self.object.agendamento.pk}
        )


class ChecklistUpdateView(
    AutomovelBaseMixin, SuccessMessageMixin,
    ChecklistFormSectionsMixin, UpdateView
):
    model = Carro_checklist
    form_class = ChecklistForm
    template_name = 'automovel/checklist_form.html'
    success_message = 'Checklist atualizado com sucesso!'

    def get_queryset(self):
        return self.apply_visibility(super().get_queryset())

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['agendamento'] = self.object.agendamento
        return context

    def get_success_url(self):
        return reverse(
            'automovel:agendamento_detail',
            kwargs={'pk': self.object.agendamento.pk}
        )


class ChecklistDetailView(AutomovelBaseMixin, DetailView):
    model = Carro_checklist
    template_name = 'automovel/checklist_detail.html'
    context_object_name = 'checklist'

    def get_queryset(self):
        return self.apply_visibility(super().get_queryset())


# ═══════════════════════════════════════════════════════════════════════════════
# MANUTENÇÃO CRUD
# ═══════════════════════════════════════════════════════════════════════════════

class ManutencaoListView(AutomovelBaseMixin, ListView):
    model = Carro_manutencao
    template_name = 'automovel/manutencao_list.html'
    context_object_name = 'manutencoes'
    paginate_by = 20

    def get_queryset(self):
        qs = super().get_queryset().select_related('carro', 'usuario')
        qs = self.apply_visibility(qs)

        carro_id = self.request.GET.get('carro')
        if carro_id:
            qs = qs.filter(carro_id=carro_id)

        status = self.request.GET.get('status')
        if status == 'concluidas':
            qs = qs.filter(concluida=True)
        elif status == 'pendentes':
            qs = qs.filter(concluida=False)

        return qs.order_by('-data_manutencao')


class ManutencaoUpdateView(AutomovelBaseMixin, UpdateView):
    model = Carro_manutencao
    form_class = ManutencaoForm
    template_name = 'automovel/manutencao_form.html'

    def get_queryset(self):
        return self.apply_visibility(super().get_queryset())

    def get_success_url(self):
        messages.success(self.request, "Manutenção atualizada com sucesso!")
        return reverse('automovel:carro_detail', kwargs={'pk': self.object.carro.pk})


class AgendarManutencaoView(AutomovelBaseMixin, View):
    apply_owner_filter = False  # Carro não tem 'usuario'

    def post(self, request, pk):
        carro_qs = Carro.objects.for_request(request)
        carro = get_object_or_404(carro_qs, pk=pk)

        form = ManutencaoForm(request.POST)
        if form.is_valid():
            manutencao = form.save(commit=False)
            manutencao.carro = carro
            manutencao.usuario = request.user
            manutencao.filial = self.get_filial_ativa()
            manutencao.save()
            messages.success(request, 'Manutenção agendada com sucesso!')
        else:
            messages.error(request, 'Erro ao agendar manutenção. Verifique os dados.')

        return redirect('automovel:carro_detail', pk=carro.pk)


# ═══════════════════════════════════════════════════════════════════════════════
# RASTREAMENTO
# ═══════════════════════════════════════════════════════════════════════════════

class RastreamentoCreateView(AutomovelBaseMixin, _ReverseGeocodeMixin, View):

    def post(self, request, *args, **kwargs):
        try:
            data = json.loads(request.body)
            base_qs = Carro_agendamento.objects.for_request(request)
            base_qs = self.apply_visibility(base_qs)
            agendamento = get_object_or_404(base_qs, pk=data.get('agendamento_id'))

            rastreamento = Carro_rastreamento.objects.create(
                agendamento=agendamento,
                latitude=data.get('latitude'),
                longitude=data.get('longitude'),
                velocidade=data.get('velocidade'),
                endereco_aproximado=self._reverse_geocode(
                    data.get('latitude'), data.get('longitude')
                ),
                filial=self.get_filial_ativa(),
            )
            return JsonResponse({
                'status': 'success',
                'id': rastreamento.id,
                'endereco': rastreamento.endereco_aproximado,
                'data_hora': rastreamento.data_hora.strftime('%d/%m/%Y %H:%M'),
            })
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=400)


class RastreamentoMapView(AutomovelBaseMixin, DetailView):
    model = Carro_agendamento
    template_name = 'automovel/rastreamento_map.html'

    def get_queryset(self):
        return self.apply_visibility(super().get_queryset())

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        rastreamentos = (
            Carro_rastreamento.objects.filter(agendamento=self.object)
            .order_by('data_hora')
        )
        context['pontos_rastreamento'] = [
            {
                'lat': float(r.latitude),
                'lng': float(r.longitude),
                'data_hora': r.data_hora.strftime('%d/%m/%Y %H:%M'),
                'velocidade': r.velocidade,
                'endereco': r.endereco_aproximado,
            }
            for r in rastreamentos
        ]
        context['mapbox_access_token'] = getattr(settings, 'MAPBOX_ACCESS_TOKEN', '')
        return context


@method_decorator(csrf_exempt, name='dispatch')
class RastreamentoAPIView(_ReverseGeocodeMixin, View):
    """Endpoint público para rastreadores físicos (sem autenticação)."""

    def post(self, request, *args, **kwargs):
        try:
            data = json.loads(request.body)
            agendamento = get_object_or_404(
                Carro_agendamento, pk=data.get('agendamento_id')
            )
            rastreamento = Carro_rastreamento.objects.create(
                agendamento=agendamento,
                latitude=data.get('latitude'),
                longitude=data.get('longitude'),
                velocidade=data.get('velocidade'),
                endereco_aproximado=self._reverse_geocode(
                    data.get('latitude'), data.get('longitude')
                ),
                filial=agendamento.filial,
            )
            return JsonResponse({'status': 'success', 'id': rastreamento.id})
        except Exception as e:
            return JsonResponse({'status': 'error', 'message': str(e)}, status=400)


# ═══════════════════════════════════════════════════════════════════════════════
# APIs
# ═══════════════════════════════════════════════════════════════════════════════

class CarrosDisponiveisAPIView(AutomovelBaseMixin, View):
    apply_owner_filter = False

    def get(self, request, *args, **kwargs):
        carros = Carro.objects.for_request(request).filter(disponivel=True, ativo=True)
        return JsonResponse(list(carros.values('id', 'placa', 'modelo')), safe=False)


class ProximaManutencaoAPIView(AutomovelBaseMixin, View):
    apply_owner_filter = False

    def get(self, request, *args, **kwargs):
        hoje = timezone.now().date()
        carros = (
            Carro.objects.for_request(request)
            .filter(data_proxima_manutencao__lte=hoje + timedelta(days=7), ativo=True)
            .order_by('data_proxima_manutencao')
        )
        data = [{'id': c.id, 'placa': c.placa, 'modelo': c.modelo} for c in carros]
        return JsonResponse(data, safe=False)


# ═══════════════════════════════════════════════════════════════════════════════
# GERADORES DE RELATÓRIO WORD
# ═══════════════════════════════════════════════════════════════════════════════

class BaseWordReportGenerator:
    def __init__(self, request, obj, filename_prefix):
        self.request = request
        self.obj = obj
        self.filename = f"{filename_prefix}_{obj.pk}_{datetime.now().strftime('%Y%m%d')}.docx"
        self.document = Document()

    def build_document(self):
        raise NotImplementedError

    def add_title(self, text):
        title = self.document.add_heading(text, level=1)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    def add_section_heading(self, text):
        self.document.add_heading(text, level=2)

    def generate(self):
        if not self.obj:
            raise Http404("Objeto não encontrado.")
        self.build_document()
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="{self.filename}"'
        self.document.save(response)
        return response


class ChecklistWordReportGenerator(BaseWordReportGenerator):
    def __init__(self, request, checklist):
        super().__init__(request, checklist, "vistoria_veicular")

    def _add_summary_table(self):
        checklist = self.obj
        ag = checklist.agendamento
        table = self.document.add_table(rows=3, cols=4)
        table.style = 'Table Grid'

        row0 = table.rows[0].cells
        row0[0].text = f"Veículo:\n{ag.carro.marca} {ag.carro.modelo} - {ag.carro.placa}"
        row0[1].text = f"Data/Hora:\n{checklist.data_hora.strftime('%d/%m/%Y %H:%M')}"
        row0[2].text = f"KM Inicial:\n{ag.km_inicial or 'N/A'} km"
        row0[3].text = f"KM Final:\n{ag.km_final or 'N/A'} km"

        row1 = table.rows[1].cells
        row1[0].text = f"Funcionário:\n{ag.funcionario or 'N/A'}"
        row1[1].text = f"Tipo:\n{checklist.get_tipo_display()}"
        row1[0].merge(row1[1])
        row1[2].merge(row1[3])

        table.rows[2].cells[0].merge(table.rows[2].cells[3])

        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def _add_items_table(self):
        checklist = self.obj
        itens = [
            ("Parte Frontal", checklist.get_revisao_frontal_status_display()),
            ("Parte Traseira", checklist.get_revisao_trazeira_status_display()),
            ("Lado do Motorista", checklist.get_revisao_lado_motorista_status_display()),
            ("Lado do Passageiro", checklist.get_revisao_lado_passageiro_status_display()),
        ]
        table = self.document.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Item Vistoriado'
        hdr[1].text = 'Status'
        for item, status in itens:
            row = table.add_row().cells
            row[0].text = item
            row[1].text = status

    def _add_photos(self):
        photos = [
            ("Evidência Frontal", self.obj.foto_frontal),
            ("Evidência Traseira", self.obj.foto_trazeira),
            ("Evidência Lado do Motorista", self.obj.foto_lado_motorista),
            ("Evidência Lado do Passageiro", self.obj.foto_lado_passageiro),
        ]
        for title, image_field in photos:
            if image_field:
                self.document.add_paragraph(title, style='Heading 3')
                try:
                    self.document.add_picture(image_field.path, width=Inches(4.0))
                except FileNotFoundError:
                    self.document.add_paragraph(f"(Imagem não encontrada: {image_field.path})")

    def build_document(self):
        self.add_title("Resumo da Vistoria do Veículo")
        self._add_summary_table()
        self.add_section_heading("Itens Vistoriados")
        self._add_items_table()
        self.add_section_heading("Observações Gerais")
        self.document.add_paragraph(self.obj.observacoes_gerais or "Nenhuma observação.")
        self.add_section_heading("Evidências Fotográficas")
        self._add_photos()


# ═══════════════════════════════════════════════════════════════════════════════
# GERADORES DE RELATÓRIO EXCEL
# ═══════════════════════════════════════════════════════════════════════════════

class BaseExcelReportGenerator:
    def __init__(self, request, queryset, filename_prefix):
        self.request = request
        self.queryset = queryset
        self.filename = f"{filename_prefix}_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
        self.header_font = Font(bold=True, color="FFFFFF")
        self.header_fill = PatternFill(start_color="4F81BD", end_color="4F81BD", fill_type="solid")
        self.header_alignment = Alignment(horizontal="center", vertical="center")
        self.title_font = Font(bold=True, size=16)
        self.title_alignment = Alignment(horizontal="center", vertical="center")
        self.thin_border = Border(
            left=Side(style='thin'), right=Side(style='thin'),
            top=Side(style='thin'), bottom=Side(style='thin'),
        )

    def get_headers(self):
        raise NotImplementedError

    def get_row_data(self, obj):
        raise NotImplementedError

    def get_report_title(self):
        raise NotImplementedError

    def _write_title(self, ws):
        last_col = get_column_letter(len(self.get_headers()))
        ws.merge_cells(f'A1:{last_col}1')
        cell = ws['A1']
        cell.value = self.get_report_title()
        cell.font = self.title_font
        cell.alignment = self.title_alignment

    def _write_headers(self, ws):
        ws.append(self.get_headers())
        for cell in ws[2]:
            cell.font = self.header_font
            cell.fill = self.header_fill
            cell.alignment = self.header_alignment
            cell.border = self.thin_border

    def _adjust_column_widths(self, ws):
        for column_cells in ws.columns:
            column_letter = column_cells[1].column_letter
            max_length = 0
            for cell in column_cells:
                if isinstance(cell, MergedCell):
                    continue
                try:
                    if cell.value and len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except Exception:
                    pass
            ws.column_dimensions[column_letter].width = max_length + 2

    def generate(self):
        response = HttpResponse(
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="{self.filename}"'
        wb = Workbook()
        ws = wb.active
        ws.title = self.get_report_title()[:30]
        self._write_title(ws)
        self._write_headers(ws)
        for obj in self.queryset:
            ws.append(self.get_row_data(obj))
        self._adjust_column_widths(ws)
        wb.save(response)
        return response


class CarroReportGenerator(BaseExcelReportGenerator):
    def __init__(self, request, queryset):
        super().__init__(request, queryset, "relatorio_carros")

    @classmethod
    def get_queryset(cls, request):
        return Carro.objects.for_request(request).filter(ativo=True).order_by('marca', 'modelo')

    def get_report_title(self):
        return "Relatório de Carros Ativos"

    def get_headers(self):
        return ['Placa', 'Marca', 'Modelo', 'Ano', 'Cor',
                'Disponível', 'Última Manutenção', 'Próxima Manutenção']

    def get_row_data(self, carro):
        return [
            carro.placa, carro.marca, carro.modelo, carro.ano, carro.cor,
            "Sim" if carro.disponivel else "Não",
            carro.data_ultima_manutencao.strftime('%d/%m/%Y') if carro.data_ultima_manutencao else 'N/A',
            carro.data_proxima_manutencao.strftime('%d/%m/%Y') if carro.data_proxima_manutencao else 'N/A',
        ]


class AgendamentoReportGenerator(BaseExcelReportGenerator):
    def __init__(self, request, queryset):
        super().__init__(request, queryset, "relatorio_agendamentos")

    @classmethod
    def get_queryset(cls, request):
        return (
            Carro_agendamento.objects.for_request(request)
            .select_related('carro').order_by('-data_hora_agenda')
        )

    def get_report_title(self):
        return "Relatório de Agendamentos"

    def get_headers(self):
        return ['ID', 'Veículo', 'Placa', 'Funcionário', 'Data Agendamento',
                'Data Devolução', 'Status', 'KM Inicial', 'KM Final', 'Descrição']

    def get_row_data(self, ag):
        return [
            ag.id,
            f"{ag.carro.marca} {ag.carro.modelo}",
            ag.carro.placa,
            str(ag.funcionario) if ag.funcionario else 'N/A',
            ag.data_hora_agenda.strftime('%d/%m/%Y %H:%M') if ag.data_hora_agenda else 'N/A',
            ag.data_hora_devolucao.strftime('%d/%m/%Y %H:%M') if ag.data_hora_devolucao else 'Pendente',
            ag.get_status_display(),
            ag.km_inicial,
            ag.km_final,
            ag.descricao,
        ]


class ChecklistReportGenerator(BaseExcelReportGenerator):
    def __init__(self, request, queryset):
        super().__init__(request, queryset, "relatorio_checklists")

    @classmethod
    def get_queryset(cls, request):
        return (
            Carro_checklist.objects.for_request(request)
            .select_related('agendamento__carro', 'usuario').order_by('-data_hora')
        )

    def get_report_title(self):
        return "Relatório de Checklists"

    def get_headers(self):
        return ['ID', 'Agendamento', 'Veículo', 'Tipo', 'Data/Hora', 'Usuário',
                'Status Frontal', 'Status Traseiro', 'Status Motorista',
                'Status Passageiro', 'Observações']

    def get_row_data(self, cl):
        return [
            cl.id,
            f"#{cl.agendamento.id}",
            cl.agendamento.carro.placa,
            cl.get_tipo_display(),
            cl.data_hora.strftime('%d/%m/%Y %H:%M'),
            cl.usuario.get_full_name(),
            cl.get_revisao_frontal_status_display(),
            cl.get_revisao_trazeira_status_display(),
            cl.get_revisao_lado_motorista_status_display(),
            cl.get_revisao_lado_passageiro_status_display(),
            cl.observacoes_gerais or 'Nenhuma',
        ]


class RastreamentoReportGenerator(BaseExcelReportGenerator):
    def __init__(self, request, queryset):
        super().__init__(request, queryset, "relatorio_rastreamento")

    @classmethod
    def get_queryset(cls, request):
        return (
            Carro_rastreamento.objects.for_request(request)
            .select_related('agendamento__carro').order_by('-data_hora')
        )

    def get_report_title(self):
        return "Relatório de Rastreamento"

    def get_headers(self):
        return ['ID', 'Agendamento', 'Veículo', 'Data/Hora',
                'Latitude', 'Longitude', 'Velocidade (km/h)', 'Endereço Aproximado']

    def get_row_data(self, r):
        return [
            r.id,
            f"#{r.agendamento.id}",
            r.agendamento.carro.placa,
            r.data_hora.strftime('%d/%m/%Y %H:%M'),
            float(r.latitude),
            float(r.longitude),
            float(r.velocidade) if r.velocidade else 'N/A',
            r.endereco_aproximado or 'N/A',
        ]


# ═══════════════════════════════════════════════════════════════════════════════
# RELATÓRIOS (Function-Based) — ÚNICAS versões (sem duplicação)
# ═══════════════════════════════════════════════════════════════════════════════

def _user_has_automovel_access(user):
    """Verifica permissão + vínculo com Funcionario."""
    if not user.is_authenticated:
        return False
    if user.is_superuser:
        return True
    if not any(p.startswith('automovel.') for p in user.get_all_permissions()):
        return False
    try:
        _ = user.funcionario
    except ObjectDoesNotExist:
        return False
    return True


@login_required
def gerar_relatorio_word(request, tipo, pk):
    if not _user_has_automovel_access(request.user):
        raise PermissionDenied("Acesso negado ao módulo Automóvel.")

    generators = {
        'checklist': (ChecklistWordReportGenerator, Carro_checklist),
    }
    config = generators.get(tipo)
    if not config:
        return HttpResponseBadRequest("Tipo de relatório inválido.")

    generator_class, model_class = config
    qs = model_class.objects.for_request(request)

    if not request.user.is_superuser and not request.user.has_perm('automovel.view_all_automovel'):
        qs = qs.filter(usuario=request.user)

    try:
        obj = qs.get(pk=pk)
    except model_class.DoesNotExist:
        raise Http404(f"{model_class._meta.verbose_name} não encontrado ou acesso negado.")

    return generator_class(request, obj).generate()


@login_required
def gerar_relatorio_excel(request, tipo):
    if not _user_has_automovel_access(request.user):
        raise PermissionDenied("Acesso negado ao módulo Automóvel.")

    generators = {
        'carros': CarroReportGenerator,
        'agendamentos': AgendamentoReportGenerator,
        'checklists': ChecklistReportGenerator,
        'rastreamento': RastreamentoReportGenerator,
    }
    generator_class = generators.get(tipo)
    if not generator_class:
        return HttpResponseBadRequest("Tipo de relatório inválido.")

    queryset = generator_class.get_queryset(request)

    # Filtra por dono somente se o model tiver campo 'usuario'
    if not request.user.is_superuser and not request.user.has_perm('automovel.view_all_automovel'):
        field_names = [f.name for f in queryset.model._meta.get_fields()]
        if 'usuario' in field_names:
            queryset = queryset.filter(usuario=request.user)

    return generator_class(request, queryset).generate()


    
