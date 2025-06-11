
from django.shortcuts import render, redirect, get_object_or_404
from .models import FichaEPI, ItemEPI, EquipamentoSeguranca, EPI
from .forms import FichaEPIForm, ItemEPIForm
from django.contrib.auth.decorators import login_required
from django.contrib import messages
from django.http import HttpResponse
from django.http import JsonResponse
from django.views.decorators.http import require_GET
from django.contrib.auth import get_user_model
from django.db.models import Q
from django.views.decorators.csrf import csrf_exempt
from django.db import transaction
from django.core.files.base import ContentFile
from django.contrib.auth.models import User

from reportlab.pdfgen import canvas
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors
from docx import Document
from docx.shared import Inches
import json
from datetime import datetime
import base64
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_CENTER


User = get_user_model()

@login_required
def buscar_empregado(request):
    matricula = request.GET.get('matricula')
    try:
        empregado = Empregado.objects.get(registro=matricula)
        data = {
            'nome': empregado.nome,
            'cargo': empregado.cargo,
            'admissao': empregado.admissao.strftime('%Y-%m-%d'),
        }
        return JsonResponse(data)
    except Empregado.DoesNotExist:
        return JsonResponse({}, status=404)

@login_required
@transaction.atomic
def criar_ficha(request):
    # Buscar todos os equipamentos ativos
    equipamentos = EquipamentoSeguranca.objects.filter(ativo=True).order_by('nome_equipamento')
    
    # Inicializar os forms fora do bloco condicional
    form_ficha = FichaEPIForm(request.POST, request.FILES)
    form_item = ItemEPIForm(request.POST)

    if request.method == 'POST':
        if form_ficha.is_valid() and form_item.is_valid():       
            try:
                # Parse dos dados JSON
                data = json.loads(request.body)
                
                # Criar ficha principal
                ficha = FichaEPI.objects.create(
                    empregado=request.user,
                    cargo=data.get('cargo', ''),
                    registro=data.get('registro', ''),
                    admissao=datetime.strptime(data['admissao'], '%Y-%m-%d').date(),
                    contrato=data.get('contrato', ''),
                    local_data=data.get('local_data', ''),
                    assinatura=save_signature(data.get('assinatura_empregado', ''))
                )
                # Processar itens de EPI
                for item in data.get('epi_items', []):
                    equipamento = EquipamentoSeguranca.objects.get(id=item['equipamento_id'])
                    
                    ItemEPI.objects.create(
                        ficha=ficha,
                        epi=equipamento,
                        quantidade=item['quantidade'],
                        data_recebimento=datetime.strptime(item['data_recebimento'], '%Y-%m-%d').date(),
                        assinatura=save_signature(item['assinatura']),
                        data_validade=calculate_expiry_date(equipamento, item['data_recebimento'])
                    )
                    
                    # Atualizar estoque
                    equipamento.quantidade_estoque -= item['quantidade']
                    equipamento.save()
                
                return JsonResponse({
                    'success': True,
                    'ficha_id': ficha.id,
                    'redirect_url': reverse('epi:visualizar_ficha', args=[ficha.id])
                })
            
            except Exception as e:
                return JsonResponse({'success': False, 'error': str(e)}, status=400)
    # Contexto para ambos GET e POST (incluindo quando o form é inválido)
    context = {
        'form_ficha': form_ficha,
        'form_item': form_item,
        'equipamentos': equipamentos
    }
    return render(request, 'epi/criar_ficha.html    ', context)

    if form_ficha.is_valid() and form_item.is_valid():
        ficha = form_ficha.save(commit=False)
        ficha.empregado = request.user
        ficha.save()
        
        # Processar os itens de EPI aqui...
        
        messages.success(request, 'Ficha de EPI criada com sucesso!')
        return redirect('epi:listar_fichas')
    else:
        form_ficha = FichaEPIForm()
        form_item = ItemEPIForm()
    
    # Passar os equipamentos para o template
    return render(request, 'epi/criar_ficha.html', {
        'form_ficha': form_ficha,
        'form_item': form_item,
        'equipamentos': equipamentos  # Esta é a chave - passar os equipamentos para o template
    })

def save_signature(signature_data):
    if not signature_data:
        return None
    
    try:
        format, imgstr = signature_data.split(';base64,') 
        ext = format.split('/')[-1] 
        file_name = f"signature_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"
        return ContentFile(base64.b64decode(imgstr), name=file_name)
    except:
        return None

def calculate_expiry_date(equipamento, receipt_date):
    if not equipamento.data_validade:
        return None
    
    receipt_date = datetime.strptime(receipt_date, '%Y-%m-%d').date()
    return min(equipamento.data_validade, receipt_date + timedelta(days=equipamento.vida_util*30))

@login_required
def listar_fichas(request):
    fichas = FichaEPI.objects.filter(empregado=request.user).prefetch_related('itens')
    return render(request, 'epi/listar_fichas.html', {'fichas': fichas})

@login_required
def visualizar_ficha(request, ficha_id):
    ficha = get_object_or_404(FichaEPI, id=ficha_id, empregado=request.user)
    itens = ficha.itens.all()
    return render(request, 'epi/visualizar_ficha.html', {'ficha': ficha, 'itens': itens})

@require_GET
@login_required
def buscar_funcionario(request):
    registro = request.GET.get('registro', '')
    
    try:
        funcionario = User.objects.get(profile__registro=registro)
        return JsonResponse({
            'nome': funcionario.get_full_name(),
            'cargo': funcionario.profile.cargo,
            'admissao': funcionario.profile.admissao.strftime('%Y-%m-%d')
        })
    except User.DoesNotExist:
        return JsonResponse({'error': 'Funcionário não encontrado'}, status=404)
    
    if not registro:
                
        data = {
            'id': funcionario.id,
            'nome': funcionario.get_full_name(),
            'registro': funcionario.profile.registro if hasattr(funcionario, 'profile') else registro,
            'cargo': funcionario.profile.cargo if hasattr(funcionario, 'profile') else '',
            'admissao': funcionario.profile.admissao.strftime('%Y-%m-%d') if hasattr(funcionario, 'profile') else '',
            'departamento': funcionario.profile.departamento if hasattr(funcionario, 'profile') else '',
        }
        
        return JsonResponse(data)
    
def gerar_pdf(request, ficha_id):
    ficha = get_object_or_404(FichaEPI, id=ficha_id)
    
    response = HttpResponse(content_type='application/pdf')
    response['Content-Disposition'] = f'attachment; filename="ficha_epi_{ficha.id}.pdf"'
    
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter)
    styles = getSampleStyleSheet()
    
    # Estilo personalizado
    styles.add(ParagraphStyle(
        name='Center',
        alignment=TA_CENTER,
        fontSize=14,
        spaceAfter=20
    ))
    
    elements = []
    
    # Cabeçalho
    elements.append(Paragraph("FICHA DE ENTREGA DE EPI", styles['Center']))
    elements.append(Paragraph(f"Empregado: {ficha.empregado.get_full_name()}", styles['Normal']))
    
    # Tabela de itens
    data = [["Equipamento", "CA", "Qtd", "Data Receb.", "Validade", "Assinatura"]]
    
    for item in ficha.itens.all():
        data.append([
            item.equipamento.nome_equipamento,
            item.equipamento.codigo_CA or '-',
            str(item.quantidade),
            item.data_recebimento.strftime('%d/%m/%Y'),
            item.data_validade.strftime('%d/%m/%Y') if item.data_validade else '-',
            "__________________"
        ])
    
    table = Table(data, colWidths=[2*inch, 1*inch, 0.5*inch, 1*inch, 1*inch, 1.5*inch])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), colors.grey),
        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
        ('FONTSIZE', (0,0), (-1,0), 10),
        ('BOTTOMPADDING', (0,0), (-1,0), 12),
        ('BACKGROUND', (0,1), (-1,-1), colors.beige),
        ('GRID', (0,0), (-1,-1), 1, colors.black),
    ]))
    
    elements.append(table)
    
    # Termo de responsabilidade
    elements.append(Paragraph("TERMO DE RESPONSABILIDADE", styles['Heading2']))
    elements.append(Paragraph(
        f"Eu, {ficha.empregado.get_full_name()}, declaro ter recebido os EPIs listados acima, "
        "comprometendo-me a utilizá-los conforme orientado e a zelar por sua conservação.",
        styles['Normal']
    ))
    
    # Assinatura
    elements.append(Paragraph("Assinatura do Empregado: ________________________", styles['Normal']))
    elements.append(Paragraph(f"Data: {timezone.now().strftime('%d/%m/%Y')}", styles['Normal']))
    
    doc.build(elements)
    pdf = buffer.getvalue()
    buffer.close()
    response.write(pdf)
    
    return response

def gerar_word(request, ficha_id):
    ficha = get_object_or_404(FichaEPI, id=ficha_id)
    
    document = Document()
    
    # Cabeçalho
    document.add_heading('FICHA DE ENTREGA DE EPI', 0)
    document.add_paragraph(f"Empregado: {ficha.empregado.get_full_name()}")
    document.add_paragraph(f"Registro: {ficha.registro}")
    document.add_paragraph(f"Data: {timezone.now().strftime('%d/%m/%Y')}")
    
    # Tabela de itens
    table = document.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    
    # Cabeçalhos
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Equipamento'
    hdr_cells[1].text = 'CA'
    hdr_cells[2].text = 'Qtd'
    hdr_cells[3].text = 'Data Receb.'
    hdr_cells[4].text = 'Validade'
    hdr_cells[5].text = 'Assinatura'
    
    # Dados
    for item in ficha.itens.all():
        row_cells = table.add_row().cells
        row_cells[0].text = item.equipamento.nome_equipamento
        row_cells[1].text = item.equipamento.codigo_CA or '-'
        row_cells[2].text = str(item.quantidade)
        row_cells[3].text = item.data_recebimento.strftime('%d/%m/%Y')
        row_cells[4].text = item.data_validade.strftime('%d/%m/%Y') if item.data_validade else '-'
        row_cells[5].text = "__________________"
    
    # Termo
    document.add_heading('Termo de Responsabilidade', level=1)
    document.add_paragraph(
        f"Eu, {ficha.empregado.get_full_name()}, declaro ter recebido os EPIs listados acima, "
        "comprometendo-me a utilizá-los conforme orientado e a zelar por sua conservação."
    )
    document.add_paragraph("Assinatura do Empregado: ________________________")
    
    # Salvar documento
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    response['Content-Disposition'] = f'attachment; filename="ficha_epi_{ficha.id}.docx"'
    
    return response

@login_required
def gerar_word(request, ficha_id):
    ficha = get_object_or_404(FichaEPI, id=ficha_id, empregado=request.user)
    itens = ficha.itens.all()

    document = Document()
    
    # Cabeçalho
    document.add_heading('FICHA DE CONTROLE DE EPI\'s', 0)
    document.add_heading('Dados de Identificação', level=2)
    
    # Dados do empregado
    dados = [
        ("Empregado:", ficha.empregado.get_full_name()),
        ("Cargo:", ficha.cargo),
        ("Registro:", ficha.registro),
        ("Admissão:", ficha.admissao.strftime('%d/%m/%Y')),
        ("Demissão:", ficha.demissao.strftime('%d/%m/%Y') if ficha.demissao else "N/A"),
        ("Contrato:", ficha.contrato),
    ]
    
    table = document.add_table(rows=0, cols=2)
    for label, value in dados:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = value
    
    # Termo de compromisso
    document.add_heading('TERMO DE COMPROMISSO', level=2)
    termo = document.add_paragraph()
    termo.add_run(f"Eu, {ficha.empregado.get_full_name()},\n")
    termo.add_run("Declaro, que recebi de CETEST MINAS ENGENHARIA E SERVIÇOS S/A, os Equipamentos de Proteção\n")
    termo.add_run("Individual - EPI's - abaixo relacionados, comprometendo-me a:\n")
    termo.add_run("1) - Usá-los em trabalho, zelando pela sua guarda e conservação, devolvendo-os quando se tornarem\n")
    termo.add_run("impróprios para o uso e/ou meu desligamento da CETEST MINAS ou do seu respectivo contrato;\n")
    termo.add_run("2) - Em caso de perda, mau uso, extravio ou inutilização proposital do EPI recebido, assumo a\n")
    termo.add_run("responsabilidade Quanto à restituição do seu valor atualizado, conforme autorização de débito por mim assinada.\n")
    termo.add_run("Declaro ainda ter recebido no ato de minha admissão e no ato do recebimento:\n")
    termo.add_run("1) Treinamento básico e instruções prévias sobre a forma de utilização e guarda dos EPI's recebido;\n")
    termo.add_run("2) Instruções sobre os riscos a que estou exposto em minha área de trabalho, bem como sua prevenção.\n")
    termo.add_run("3) Estou ciente de que o não uso dos EPI's, constitui ato faltoso conforme artigo 158 da CLT.\n\n")
    termo.add_run(f"Local e Data: {ficha.local_data}\n")
    termo.add_run("Assinatura: ___________________________________________\n")
    
    # Tabela de EPIs
    document.add_heading('RECEBIMENTO / DEVOLUÇÃO', level=2)
    
    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Data'
    hdr_cells[1].text = 'Qtde.'
    hdr_cells[2].text = 'Un.'
    hdr_cells[3].text = 'Descrição EPI'
    hdr_cells[4].text = 'Certificado Aprovação'
    hdr_cells[5].text = 'Assinatura'
    
    for item in itens:
        row_cells = table.add_row().cells
        row_cells[0].text = item.data_recebimento.strftime('%d/%m/%Y')
        row_cells[1].text = str(item.quantidade)
        row_cells[2].text = item.epi.unidade
        row_cells[3].text = item.epi.nome
        row_cells[4].text = item.epi.certificado
        row_cells[5].text = "_________________________"
    
    # Artigo CLT
    document.add_heading('CONSOLIDAÇÃO DAS LEIS DO TRABALHO', level=2)
    document.add_paragraph('Art. 158 - Cabe aos empregados', style='ListBullet')
    document.add_paragraph('I - Observar as normas de segurança e medicina do trabalho, inclusive as instruções de que trata o item II do artigo anterior;', style='ListBullet2')
    document.add_paragraph('II-Colaborar com a empresa na aplicação dos dispositivos deste capítulo.', style='ListBullet2')
    document.add_paragraph('Parágrafo Único: Constitui ato faltoso do empregado a recusa injustificada:', style='ListBullet2')
    document.add_paragraph('a) ...', style='ListBullet3')
    document.add_paragraph('b) ao uso dos equipamentos de proteção individual fornecidos pela empresa.', style='ListBullet3')
    
    # Salvar o documento
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    
    response = HttpResponse(buffer, content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="ficha_epi_{ficha.empregado.username}.docx"'
    
    return response
