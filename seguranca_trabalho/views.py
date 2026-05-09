
# seguranca_trabalho/views.py

import base64
import io
import json
import logging
from datetime import timedelta
from pathlib import Path
from django.views.generic.detail import SingleObjectMixin
from celery.exceptions import ImproperlyConfigured
from django.conf import settings
from django.contrib import messages
from django.contrib.auth.decorators import login_required, permission_required
from django.core.exceptions import PermissionDenied
from django.db import IntegrityError, transaction
from django.db.models import Q, Count, Sum, Value, IntegerField, ProtectedError
from django.db.models.functions import Coalesce
from django.http import HttpResponse
from django.shortcuts import get_object_or_404, redirect, render
from django.template.loader import render_to_string
from django.urls import reverse, reverse_lazy
from django.utils import timezone
from django.utils.safestring import mark_safe
from django.utils.translation import gettext_lazy as _
from django.views import View
from django.views.generic import (
    CreateView, DeleteView, DetailView, ListView, TemplateView, UpdateView,
)
from django.views.generic.edit import FormMixin
from django.views.decorators.http import require_POST

from docx import Document
from weasyprint import HTML, default_url_fetcher

from core.mixins import (
    AppPermissionMixin, FuncionarioRequiredMixin, ViewFilialScopedMixin,
    TecnicoScopeMixin, FilialCreateMixin, LoginRequiredMixin,
)
from departamento_pessoal.models import Funcionario
from usuario.models import Filial

from .forms import (
    AssinaturaEntregaForm, AssinaturaTermoForm, EntregaEPIForm,
    EquipamentoForm, FichaEPIForm, FuncaoForm, CargoFuncaoForm,
    AjusteEstoqueForm,
)
from .models import (
    EntregaEPI, Equipamento, FichaEPI, Funcao, MatrizEPI,
    CargoFuncao, MovimentacaoEstoque,
)


logger = logging.getLogger(__name__)

_APP = 'seguranca_trabalho'


# =============================================================================
# HELPERS
# =============================================================================

def custom_url_fetcher(url):
    """Permite que o WeasyPrint acesse arquivos de media locais."""
    if url.startswith(settings.MEDIA_URL):
        path = (settings.MEDIA_ROOT / url[len(settings.MEDIA_URL):]).as_posix()
        return default_url_fetcher(f'file://{path}')
    return default_url_fetcher(url)


def _estoque_equipamento(equipamento, filial):
    """Calcula estoque atual de um equipamento na filial."""
    mov = MovimentacaoEstoque.objects.filter(equipamento=equipamento, filial=filial)
    entradas = mov.filter(tipo='ENTRADA').aggregate(
        total=Coalesce(Sum('quantidade'), Value(0, output_field=IntegerField()))
    )['total']
    saidas = mov.filter(tipo='SAIDA').aggregate(
        total=Coalesce(Sum('quantidade'), Value(0, output_field=IntegerField()))
    )['total']
    return entradas - saidas


def _processar_assinatura_base64(sig_str):
    """Normaliza uma string de assinatura base64 para data URI válido."""
    if not sig_str:
        return None
    sig = sig_str.strip()
    if sig.startswith('data:image'):
        return mark_safe(sig)
    if len(sig) > 100:
        return mark_safe(f'data:image/png;base64,{sig}')
    return None


def _imagem_file_para_base64(image_field):
    """Converte um ImageField/FileField em data URI base64."""
    if not image_field or not image_field.name:
        return None
    try:
        image_field.open('rb')
        content = image_field.read()
        image_field.close()
        if not content:
            return None
        mime = 'image/jpeg' if content[:2] == b'\xff\xd8' else 'image/png'
        encoded = base64.b64encode(content).decode('utf-8')
        return mark_safe(f'data:{mime};base64,{encoded}')
    except Exception as e:
        logger.warning("Erro ao converter imagem para base64: %s", e)
        return None


# =============================================================================
# MIXINS DE SUPORTE
# =============================================================================

class FilialAtivaMixin:
    """
    Obtém a filial ativa com prioridade para o seletor da sessão.
    Fallbacks: user.filial_ativa → funcionario.filial.
    """

    def get_filial_ativa(self):
        filial_id = self.request.session.get('active_filial_id')
        if filial_id:
            try:
                return Filial.objects.get(pk=filial_id)
            except Filial.DoesNotExist:
                pass

        user = self.request.user
        filial_ativa = getattr(user, 'filial_ativa', None)
        if filial_ativa:
            return filial_ativa

        try:
            funcionario = Funcionario.objects.select_related('filial').get(usuario=user)
            return funcionario.filial
        except Funcionario.DoesNotExist:
            return None

    def get_filial_ativa_id(self):
        filial = self.get_filial_ativa()
        return filial.id if filial else None


class SSTVisibilityMixin(FilialAtivaMixin):
    """
    Controla a visibilidade dos registros de SST conforme o perfil do usuário.

    Regras:
      - Superuser → tudo
      - Permissão global 'view_all_*' → tudo da filial (já filtrada)
      - Demais usuários com Funcionario → registros da filial dele (FilialManager)
      - Sem vínculo → nada
    """

    def get_queryset(self):
        parent = super()
        if hasattr(parent, 'get_queryset'):
            qs = parent.get_queryset()
        elif self.model is not None:
            qs = self.model._default_manager.all()
        else:
            raise ImproperlyConfigured(
                f"{self.__class__.__name__} precisa definir 'model' "
                f"ou herdar de uma view com get_queryset()."
            )
        # ... resto da lógica de filtro por filial
        filial = self.get_filial_ativa()
        if filial and hasattr(qs.model, 'filial'):
            qs = qs.filter(filial=filial)
        return qs

    def apply_visibility(self, queryset):
        user = self.request.user

        if user.is_superuser:
            return queryset

        if user.has_perm('seguranca_trabalho.view_all_seguranca_trabalho'):
            return queryset

        funcionario = getattr(user, 'funcionario', None)
        if not funcionario:
            return queryset.none()

        return queryset


class SSTSearchMixin:
    """
    Aplica busca textual via parâmetro `?q=` em campos definidos por
    `search_fields`. Inclui ordering opcional via `search_order_by`.

    Exemplo:
        class MinhaView(SSTSearchMixin, ListView):
            search_fields = ['funcionario__nome_completo', 'funcionario__matricula']
            search_order_by = 'funcionario__nome_completo'
    """
    search_fields: list[str] = []
    search_order_by: str | None = None
    search_distinct: bool = False

    def apply_search(self, queryset):
        q = self.request.GET.get('q')
        if q and self.search_fields:
            filters = Q()
            for field in self.search_fields:
                filters |= Q(**{f'{field}__icontains': q})
            queryset = queryset.filter(filters)
            if self.search_distinct:
                queryset = queryset.distinct()
        if self.search_order_by:
            queryset = queryset.order_by(self.search_order_by)
        return queryset


# =============================================================================
# MIXIN BASE DO APP
# =============================================================================

class SSTBaseMixin(
    FuncionarioRequiredMixin,
    AppPermissionMixin,
    SSTVisibilityMixin,
    ViewFilialScopedMixin,
):
    """
    Mixin base para todas as CBVs do app Segurança do Trabalho.

    MRO:
      1. FuncionarioRequiredMixin → autenticação + vínculo com Funcionario
      2. AppPermissionMixin       → permissão por app ('seguranca_trabalho')
      3. SSTVisibilityMixin       → apply_visibility() + FilialAtivaMixin
      4. ViewFilialScopedMixin    → get_queryset() filtrado por filial

    O `get_queryset()` aqui já aplica visibilidade automaticamente,
    portanto subclasses NÃO precisam chamar `apply_visibility()` manualmente.

    Para escopo técnico, herdar também de `TecnicoScopeMixin`
    e definir `tecnico_scope_lookup`.
    """
    modulo_nome = 'Segurança do Trabalho'
    app_label_required = _APP

    def dispatch(self, request, *args, **kwargs):
        if request.user.is_authenticated and not request.user.is_superuser:
            if not request.user.has_perm(f'{_APP}.view_all_{_APP}'):
                try:
                    _ = request.user.funcionario
                except Funcionario.DoesNotExist:
                    return render(request, 'seguranca_trabalho/acesso_negado.html', {
                        'titulo': 'Acesso Restrito',
                        'mensagem': (
                            'Sua conta não está vinculada a um registro de '
                            'funcionário, por isso não pode acessar o módulo '
                            'de Segurança do Trabalho.'
                        ),
                    }, status=403)
        return super().dispatch(request, *args, **kwargs)

    def get_queryset(self):
        """Aplica filial + visibilidade de uma vez para evitar duplicação."""
        qs = super().get_queryset()
        return self.apply_visibility(qs)


# =============================================================================
# EQUIPAMENTOS (CRUD + Ajuste de Estoque)
# =============================================================================

class EquipamentoListView(SSTBaseMixin, ListView):
    model = Equipamento
    permission_required = 'seguranca_trabalho.view_equipamento'
    template_name = 'seguranca_trabalho/equipamento_list.html'
    context_object_name = 'equipamentos'
    paginate_by = 20

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        filial = self.get_filial_ativa()

        if filial:
            # Pré-calcula estoque em UMA query
            movs_por_equip = (
                MovimentacaoEstoque.objects
                .filter(filial=filial)
                .values('equipamento_id', 'tipo')
                .annotate(total=Sum('quantidade'))
            )

            mapa = {}
            for item in movs_por_equip:
                eq_id = item['equipamento_id']
                mapa.setdefault(eq_id, {'ENTRADA': 0, 'SAIDA': 0})
                mapa[eq_id][item['tipo']] = item['total'] or 0

            for eq in context['equipamentos']:
                dados = mapa.get(eq.pk, {'ENTRADA': 0, 'SAIDA': 0})
                eq.total_entradas = dados['ENTRADA']
                eq.total_saidas = dados['SAIDA']
                eq.estoque_atual = dados['ENTRADA'] - dados['SAIDA']

        return context


class EquipamentoDetailView(SSTBaseMixin, DetailView):
    model = Equipamento
    permission_required = 'seguranca_trabalho.view_equipamento'
    template_name = 'seguranca_trabalho/equipamento_detail.html'
    context_object_name = 'equipamento'


class EquipamentoCreateView(SSTBaseMixin, FilialCreateMixin, CreateView):
    model = Equipamento
    permission_required = 'seguranca_trabalho.add_equipamento'
    form_class = EquipamentoForm
    template_name = 'seguranca_trabalho/equipamento_form.html'
    success_url = reverse_lazy('seguranca_trabalho:equipamento_list')

    def form_valid(self, form):
        response = super().form_valid(form)

        estoque_inicial = form.cleaned_data.get('estoque_inicial')
        if estoque_inicial and estoque_inicial > 0:
            MovimentacaoEstoque.objects.create(
                equipamento=self.object,
                tipo='ENTRADA',
                quantidade=estoque_inicial,
                justificativa='Carga inicial de estoque (cadastro do equipamento)',
                responsavel=self.request.user,
                filial=self.object.filial,
            )
        return response


class EquipamentoUpdateView(SSTBaseMixin, UpdateView):
    model = Equipamento
    permission_required = 'seguranca_trabalho.change_equipamento'
    form_class = EquipamentoForm
    template_name = 'seguranca_trabalho/equipamento_form.html'
    success_url = reverse_lazy('seguranca_trabalho:equipamento_list')

    def form_valid(self, form):
        messages.success(
            self.request,
            f"Equipamento '{form.instance.nome}' atualizado com sucesso!"
        )
        return super().form_valid(form)


class EquipamentoDeleteView(SSTBaseMixin, DeleteView):
    model = Equipamento
    permission_required = 'seguranca_trabalho.delete_equipamento'
    template_name = 'seguranca_trabalho/confirm_delete.html'
    success_url = reverse_lazy('seguranca_trabalho:equipamento_list')
    context_object_name = 'object'

    def form_valid(self, form):
        messages.success(self.request, "Equipamento excluído com sucesso!")
        return super().form_valid(form)


class AjusteEstoqueView(SSTBaseMixin, View):
    """Permite ajustar o estoque de um equipamento com justificativa."""
    http_method_names = ['get', 'post']

    def _get_equipamento(self, request, pk):
        filial = self.get_filial_ativa()
        return get_object_or_404(Equipamento, pk=pk, filial=filial)

    def _render(self, request, equipamento, form):
        return render(request, 'seguranca_trabalho/ajuste_estoque.html', {
            'equipamento': equipamento,
            'form': form,
        })

    def get(self, request, pk):
        equipamento = self._get_equipamento(request, pk)
        return self._render(request, equipamento, AjusteEstoqueForm())

    def post(self, request, pk):
        equipamento = self._get_equipamento(request, pk)
        filial = self.get_filial_ativa()
        form = AjusteEstoqueForm(request.POST)

        if not form.is_valid():
            return self._render(request, equipamento, form)

        tipo = form.cleaned_data['tipo']
        quantidade = form.cleaned_data['quantidade']
        justificativa = form.cleaned_data['justificativa']

        estoque_atual = _estoque_equipamento(equipamento, filial)

        if tipo == 'SAIDA' and quantidade > estoque_atual:
            form.add_error(
                'quantidade',
                _(f"Estoque insuficiente. Disponível: {estoque_atual}")
            )
            return self._render(request, equipamento, form)

        MovimentacaoEstoque.objects.create(
            equipamento=equipamento,
            tipo=tipo,
            quantidade=quantidade,
            justificativa=f"[AJUSTE MANUAL] {justificativa}",
            responsavel=request.user,
            filial=filial,
        )

        tipo_label = "adicionadas ao" if tipo == 'ENTRADA' else "removidas do"
        messages.success(
            request,
            f"{quantidade} unidades {tipo_label} estoque de '{equipamento.nome}'."
        )
        return redirect('seguranca_trabalho:equipamento_detail', pk=equipamento.pk)


# =============================================================================
# FICHAS EPI (CRUD)
# =============================================================================

class FichaEPIListView(SSTBaseMixin, TecnicoScopeMixin, SSTSearchMixin, ListView):
    model = FichaEPI
    permission_required = 'seguranca_trabalho.view_fichaepi'
    template_name = 'seguranca_trabalho/ficha_list.html'
    context_object_name = 'fichas'
    paginate_by = 30
    tecnico_scope_lookup = 'funcionario__usuario'
    search_fields = ['funcionario__nome_completo', 'funcionario__matricula']
    search_order_by = 'funcionario__nome_completo'

    def get_queryset(self):
        qs = super().get_queryset().select_related('funcionario', 'funcionario__cargo')
        return self.apply_search(qs)


class FichaEPICreateView(SSTBaseMixin, CreateView):
    model = FichaEPI
    permission_required = 'seguranca_trabalho.add_fichaepi'
    form_class = FichaEPIForm
    template_name = 'seguranca_trabalho/ficha_create.html'
    success_url = reverse_lazy('seguranca_trabalho:ficha_list')

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['request'] = self.request
        return kwargs

    def form_valid(self, form):
        # A filial da ficha é derivada do funcionário selecionado.
        form.instance.filial = form.cleaned_data['funcionario'].filial
        return super().form_valid(form)


class FichaEPIDetailView(SSTBaseMixin, TecnicoScopeMixin, FormMixin, DetailView):
    model = FichaEPI
    permission_required = 'seguranca_trabalho.view_fichaepi'
    template_name = 'seguranca_trabalho/ficha_detail.html'
    context_object_name = 'ficha'
    form_class = EntregaEPIForm
    tecnico_scope_lookup = 'funcionario__usuario'

    def get_queryset(self):
        return super().get_queryset().select_related(
            'funcionario', 'funcionario__cargo'
        )

    def get_success_url(self):
        return reverse('seguranca_trabalho:ficha_detail', kwargs={'pk': self.object.pk})

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['filial'] = self.get_object().filial
        return kwargs

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['form'] = self.get_form()
        context['entregas'] = (
            self.object.entregas
            .select_related('equipamento')
            .order_by('-data_entrega')
        )
        return context

    def post(self, request, *args, **kwargs):
        if not request.user.has_perm('seguranca_trabalho.add_entregaepi'):
            raise PermissionDenied("Você não tem permissão para registrar uma nova entrega.")
        self.object = self.get_object()
        form = self.get_form()
        if form.is_valid():
            return self.form_valid(form)
        return self.form_invalid(form)

    def form_valid(self, form):
        with transaction.atomic():
            nova_entrega = form.save(commit=False)
            nova_entrega.ficha = self.object
            nova_entrega.filial = self.object.filial
            nova_entrega.save()

            MovimentacaoEstoque.objects.create(
                equipamento=nova_entrega.equipamento,
                tipo='SAIDA',
                quantidade=nova_entrega.quantidade,
                responsavel=self.request.user,
                justificativa=(
                    f"Entrega EPI - Ficha #{self.object.pk} "
                    f"({self.object.funcionario.nome_completo})"
                ),
                entrega_associada=nova_entrega,
                filial=nova_entrega.filial,
            )
        messages.success(self.request, "Nova entrega de EPI registrada com sucesso!")
        return redirect(self.get_success_url())


class FichaEPIUpdateView(SSTBaseMixin, UpdateView):
    model = FichaEPI
    permission_required = 'seguranca_trabalho.change_fichaepi'
    form_class = FichaEPIForm
    template_name = 'seguranca_trabalho/ficha_form.html'

    def get_success_url(self):
        return reverse('seguranca_trabalho:ficha_detail', kwargs={'pk': self.object.pk})

    def get_form_kwargs(self):
        kwargs = super().get_form_kwargs()
        kwargs['request'] = self.request
        return kwargs


class FichaEPIDeleteView(SSTBaseMixin, DeleteView):
    model = FichaEPI
    permission_required = 'seguranca_trabalho.delete_fichaepi'
    template_name = 'seguranca_trabalho/confirm_delete.html'
    success_url = reverse_lazy('seguranca_trabalho:ficha_list')
    context_object_name = 'object'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        ficha = self.object
        entregas = ficha.entregas.select_related('equipamento')
        context['has_entregas'] = entregas.exists()
        context['entregas_count'] = entregas.count()
        return context

    def form_valid(self, form):
        try:
            messages.success(self.request, "Ficha de EPI excluída com sucesso!")
            return super().form_valid(form)
        except ProtectedError:
            ficha = self.get_object()
            count = ficha.entregas.count()
            messages.error(
                self.request,
                f"Não é possível excluir esta ficha. "
                f"Existem {count} entrega(s) de EPI vinculada(s). "
                f"Remova ou devolva todas as entregas antes de excluir a ficha."
            )
            return redirect('seguranca_trabalho:ficha_detail', pk=ficha.pk)


# =============================================================================
# ENTREGAS EPI (Assinatura e Devolução)
# =============================================================================

class AssinarEntregaView(SSTBaseMixin, TecnicoScopeMixin, UpdateView):
    model = EntregaEPI
    permission_required = 'seguranca_trabalho.change_entregaepi'
    http_method_names = ['get', 'post']
    form_class = AssinaturaEntregaForm
    template_name = 'seguranca_trabalho/entrega_sign.html'
    context_object_name = 'entrega'
    tecnico_scope_lookup = 'ficha__funcionario__usuario'

    def get_queryset(self):
        return super().get_queryset().select_related(
            'ficha', 'ficha__funcionario', 'equipamento'
        )

    def get_success_url(self):
        return reverse(
            'seguranca_trabalho:ficha_detail',
            kwargs={'pk': self.object.ficha.pk}
        )

    def get(self, request, *args, **kwargs):
        self.object = self.get_object()
        if self.object.data_devolucao or self.object.data_assinatura:
            messages.info(request, "Esta entrega já foi processada.")
            return redirect(self.get_success_url())
        return super().get(request, *args, **kwargs)

    def form_valid(self, form):
        entrega = form.save(commit=False)

        assinatura_base64 = self.request.POST.get('assinatura_base64', '').strip()
        if assinatura_base64:
            entrega.assinatura_recebimento = assinatura_base64

        if 'assinatura_imagem' in self.request.FILES:
            entrega.assinatura_imagem = self.request.FILES['assinatura_imagem']

        entrega.data_assinatura = timezone.now()
        entrega.save()

        messages.success(self.request, "Assinatura registrada com sucesso!")
        return redirect(self.get_success_url())


class RegistrarDevolucaoView(SSTBaseMixin, LoginRequiredMixin, SingleObjectMixin, View):
    model = EntregaEPI
    pk_url_kwarg = 'pk'
    app_label_required = 'seguranca_trabalho'  # ajuste se o app_label for outro

    # GET não faz sentido aqui — só aceita POST
    def get(self, request, *args, **kwargs):
        return redirect('seguranca_trabalho:entrega_list')

    def post(self, request, *args, **kwargs):
        entrega = self.get_object()  # já respeita o queryset com escopo de filial

        if entrega.data_devolucao:
            messages.warning(request, "Esta entrega já foi devolvida.")
            return redirect('seguranca_trabalho:entrega_list')

        entrega.data_devolucao = timezone.now()
        entrega.save(update_fields=['data_devolucao'])
        messages.success(request, "Devolução registrada com sucesso.")
        return redirect('seguranca_trabalho:entrega_list')


# =============================================================================
# RELATÓRIOS PDF - FICHA INDIVIDUAL
# =============================================================================

class GerarFichaPDFView(SSTBaseMixin, TecnicoScopeMixin, DetailView):
    model = FichaEPI
    permission_required = 'seguranca_trabalho.view_fichaepi'
    template_name = 'seguranca_trabalho/ficha_pdf_template.html'
    context_object_name = 'ficha'
    tecnico_scope_lookup = 'funcionario__usuario'

    def get_queryset(self):
        return super().get_queryset().select_related(
            'funcionario', 'funcionario__cargo', 'filial'
        )

    @staticmethod
    def _processar_assinatura(entrega):
        resultado = _processar_assinatura_base64(entrega.assinatura_recebimento)
        if resultado:
            return resultado
        return _imagem_file_para_base64(entrega.assinatura_imagem)

    def _get_logo_base64(self, filial):
        if filial and hasattr(filial, 'logo'):
            logo = _imagem_file_para_base64(filial.logo)
            if logo:
                return logo

        for nome in ['logo.png', 'logo.jpg', 'logo_cetest.png']:
            logo_path = Path(settings.BASE_DIR) / 'static' / 'images' / nome
            if logo_path.exists():
                content = logo_path.read_bytes()
                mime = 'image/jpeg' if content[:2] == b'\xff\xd8' else 'image/png'
                encoded = base64.b64encode(content).decode('utf-8')
                return mark_safe(f'data:{mime};base64,{encoded}')

        return None

    def get(self, request, *args, **kwargs):
        ficha = self.get_object()
        logger.info("Gerando ficha PDF: %s", ficha.funcionario.nome_completo)

        entregas = (
            EntregaEPI.objects
            .filter(ficha=ficha)
            .select_related('equipamento')
            .order_by('data_entrega')
        )

        for entrega in entregas:
            entrega.assinatura_base64 = self._processar_assinatura(entrega)

        context = {
            'ficha': ficha,
            'entregas': entregas,
            'data_emissao': timezone.now(),
            'logo_base64': self._get_logo_base64(ficha.filial),
            'assinatura_funcionario': _processar_assinatura_base64(
                ficha.assinatura_funcionario
            ),
        }

        html_string = render_to_string(self.template_name, context)

        if settings.DEBUG:
            debug_path = Path(settings.BASE_DIR) / 'debug_ficha.html'
            debug_path.write_text(html_string, encoding='utf-8')
            logger.debug("Debug HTML: %s", debug_path)

        html = HTML(
            string=html_string,
            base_url=request.build_absolute_uri(),
            url_fetcher=custom_url_fetcher,
        )
        pdf = html.write_pdf()
        logger.info("PDF gerado (%d bytes)", len(pdf))

        response = HttpResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = (
            f'attachment; filename="ficha_epi_{ficha.funcionario.matricula}.pdf"'
        )
        return response


# =============================================================================
# TERMOS E ASSINATURAS
# =============================================================================

class AssinarTermoView(SSTBaseMixin, TecnicoScopeMixin, UpdateView):
    model = FichaEPI
    permission_required = 'seguranca_trabalho.change_fichaepi'
    http_method_names = ['get', 'post']
    form_class = AssinaturaTermoForm
    template_name = 'seguranca_trabalho/termo_sign.html'
    context_object_name = 'ficha'
    tecnico_scope_lookup = 'funcionario__usuario'

    def get_queryset(self):
        return super().get_queryset().select_related(
            'funcionario', 'funcionario__cargo'
        )

    def get_success_url(self):
        return reverse('seguranca_trabalho:ficha_detail', kwargs={'pk': self.object.pk})

    def get(self, request, *args, **kwargs):
        self.object = self.get_object()
        if self.object.assinatura_funcionario:
            messages.info(request, "O termo já foi assinado pelo funcionário.")
            return redirect(self.get_success_url())
        return super().get(request, *args, **kwargs)

    def form_valid(self, form):
        ficha = form.save(commit=False)

        assinatura_base64 = self.request.POST.get('assinatura_base64', '').strip()
        if not assinatura_base64:
            messages.error(self.request, "Por favor, assine no campo de assinatura.")
            return self.form_invalid(form)

        ficha.assinatura_funcionario = assinatura_base64
        ficha.data_assinatura_termo = timezone.now()
        ficha.save()

        messages.success(self.request, "Termo assinado com sucesso!")
        return redirect(self.get_success_url())


# =============================================================================
# DASHBOARD SST
# =============================================================================

class DashboardSSTView(SSTBaseMixin, TecnicoScopeMixin, TemplateView):
    template_name = 'seguranca_trabalho/dashboard.html'
    tecnico_scope_lookup = 'funcionario__usuario'
    permission_required = 'seguranca_trabalho.view_dashboard'

    def _is_tecnico(self):
        return getattr(self.request.user, 'is_tecnico', False)

    def _qs_filial(self, model_class):
        """Retorna queryset do model filtrado pela filial ativa."""
        filial_id = self.get_filial_ativa_id()
        if not filial_id:
            return model_class.objects.none()

        if hasattr(model_class, 'filial'):
            return model_class.objects.filter(filial_id=filial_id)
        if hasattr(model_class, 'funcionario'):
            return model_class.objects.filter(funcionario__filial_id=filial_id)
        return model_class.objects.none()

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)

        try:
            equipamentos = self._qs_filial(Equipamento)
            fichas = self._qs_filial(FichaEPI)
            entregas = self._qs_filial(EntregaEPI)
            matriz = self._qs_filial(MatrizEPI)
        except Exception as e:
            logger.error("Erro no DashboardSST: %s", e)
            equipamentos = Equipamento.objects.none()
            fichas = FichaEPI.objects.none()
            entregas = EntregaEPI.objects.none()
            matriz = MatrizEPI.objects.none()

        # Técnicos só veem os próprios dados
        if self._is_tecnico():
            equipamentos = equipamentos.none()
            matriz = matriz.none()
            fichas = fichas.filter(funcionario__usuario=self.request.user)
            entregas = entregas.filter(ficha__funcionario__usuario=self.request.user)

        # ---------- KPIs ----------
        context['total_equipamentos_ativos'] = equipamentos.filter(ativo=True).count()
        context['fichas_ativas'] = fichas.filter(funcionario__status='ATIVO').count()
        context['entregas_pendentes_assinatura'] = entregas.filter(
            data_devolucao__isnull=True,
            data_assinatura__isnull=True
        ).count()

        # ---------- GRÁFICO: Status de Vencimento ----------
        today = timezone.now().date()
        thirty_days = today + timedelta(days=30)
        entregas_ativas = entregas.filter(
            data_devolucao__isnull=True,
            data_entrega__isnull=False
        ).select_related('equipamento')

        epis_vencidos = epis_vencendo = epis_regulares = 0
        for entrega in entregas_ativas:
            if not entrega.equipamento.vida_util_dias:
                epis_regulares += 1
                continue
            vencimento = entrega.data_entrega + timedelta(days=entrega.equipamento.vida_util_dias)
            if vencimento < today:
                epis_vencidos += 1
            elif today <= vencimento <= thirty_days:
                epis_vencendo += 1
            else:
                epis_regulares += 1

        context['epis_vencendo_em_30_dias'] = epis_vencendo
        context['chart_vencimento_labels'] = json.dumps(['Regulares', 'Vencendo (30d)', 'Vencidos'])
        context['chart_vencimento_data'] = json.dumps([epis_regulares, epis_vencendo, epis_vencidos])

        # ---------- GRÁFICO: Matriz de EPI ----------
        matriz_data = (
            matriz.values('funcao__nome')
            .annotate(num_epis=Count('equipamento'))
            .order_by('-num_epis')[:10]
        )
        if matriz_data:
            context['matriz_labels'] = json.dumps([m['funcao__nome'] for m in matriz_data])
            context['matriz_data'] = json.dumps([m['num_epis'] for m in matriz_data])

        # ---------- GRÁFICO: Status das Entregas ----------
        entregas_assinadas = entregas.filter(
            data_devolucao__isnull=True,
            data_assinatura__isnull=False
        ).count()
        entregas_pendentes = context['entregas_pendentes_assinatura']
        entregas_devolvidas = entregas.filter(data_devolucao__isnull=False).count()

        context['chart_status_entregas_labels'] = json.dumps(
            ['Assinadas (Ativas)', 'Pendentes', 'Devolvidas']
        )
        context['chart_status_entregas_data'] = json.dumps(
            [entregas_assinadas, entregas_pendentes, entregas_devolvidas]
        )

        # ---------- GRÁFICO: Top 5 EPIs ----------
        top_epis = (
            entregas.values('equipamento__nome')
            .annotate(total=Count('id'))
            .order_by('-total')[:5]
        )
        if top_epis:
            context['chart_top_epis_labels'] = json.dumps([e['equipamento__nome'] for e in top_epis])
            context['chart_top_epis_data'] = json.dumps([e['total'] for e in top_epis])

        context['titulo_pagina'] = "Painel de Segurança do Trabalho"
        return context


# =============================================================================
# FUNÇÕES (Cargos do Trabalho)
# =============================================================================

class FuncaoListView(SSTBaseMixin, SSTSearchMixin, ListView):
    model = Funcao
    permission_required = 'seguranca_trabalho.view_funcao'
    template_name = 'seguranca_trabalho/funcao_list.html'
    context_object_name = 'funcoes'
    paginate_by = 15
    search_fields = ['nome', 'descricao', 'funcoes_cargo__cargo__nome']
    search_order_by = 'nome'
    search_distinct = True

    def get_queryset(self):
        qs = super().get_queryset().prefetch_related('funcoes_cargo__cargo')
        return self.apply_search(qs)


class FuncaoCreateView(SSTBaseMixin, FilialCreateMixin, CreateView):
    model = Funcao
    permission_required = 'seguranca_trabalho.add_funcao'
    form_class = FuncaoForm
    success_url = reverse_lazy('seguranca_trabalho:funcao_list')

    def get_template_names(self):
        if self.request.htmx:
            return ['seguranca_trabalho/partials/base_modal.html']
        return ['seguranca_trabalho/funcao_form.html']


class FuncaoUpdateView(SSTBaseMixin, UpdateView):
    model = Funcao
    permission_required = 'seguranca_trabalho.change_funcao'
    form_class = FuncaoForm
    success_url = reverse_lazy('seguranca_trabalho:funcao_list')

    def form_valid(self, form):
        messages.success(
            self.request,
            f"Função '{form.instance.nome}' atualizada com sucesso."
        )
        return super().form_valid(form)

    def get_template_names(self):
        if self.request.htmx:
            return ['seguranca_trabalho/partials/base_modal.html']
        return ['seguranca_trabalho/funcao_form.html']


class FuncaoDeleteView(SSTBaseMixin, DeleteView):
    model = Funcao
    permission_required = 'seguranca_trabalho.delete_funcao'
    template_name = 'seguranca_trabalho/funcao_confirm_delete.html'
    success_url = reverse_lazy('seguranca_trabalho:funcao_list')
    context_object_name = 'object'

    def form_valid(self, form):
        messages.success(self.request, "Função excluída com sucesso.")
        return super().form_valid(form)


# =============================================================================
# ASSOCIAÇÕES (Função x Cargo x EPI)
# =============================================================================

class AssociacaoListView(SSTBaseMixin, SSTSearchMixin, ListView):
    model = CargoFuncao
    permission_required = 'seguranca_trabalho.view_cargofuncao'
    context_object_name = 'associacoes'
    template_name = 'seguranca_trabalho/lista_associacoes.html'
    paginate_by = 20
    search_fields = ['cargo__nome', 'funcao__nome']
    search_order_by = 'cargo__nome'

    def get_queryset(self):
        qs = super().get_queryset().select_related('cargo', 'funcao')
        return self.apply_search(qs)


class AssociacaoCreateView(SSTBaseMixin, FilialCreateMixin, CreateView):
    model = CargoFuncao
    permission_required = 'seguranca_trabalho.add_cargofuncao'
    form_class = CargoFuncaoForm
    template_name = 'seguranca_trabalho/formulario_associacao.html'
    success_url = reverse_lazy('seguranca_trabalho:lista_associacoes')

    def form_valid(self, form):
        try:
            return super().form_valid(form)
        except IntegrityError:
            cargo = form.cleaned_data.get('cargo')
            funcao = form.cleaned_data.get('funcao')
            messages.error(
                self.request,
                f"A associação entre o cargo '{cargo}' e a função '{funcao}' já existe."
            )
            return self.form_invalid(form)


@login_required
@permission_required('seguranca_trabalho.delete_cargofuncao', raise_exception=True)
@require_POST
def desvincular_funcao_cargo(request, funcao_id, cargo_id):
    """Remove o vínculo entre um cargo e uma função via HTMX."""
    filial_id = request.session.get('active_filial_id')
    associacao = get_object_or_404(
        CargoFuncao,
        funcao_id=funcao_id,
        cargo_id=cargo_id,
        funcao__filial_id=filial_id
    )
    associacao.delete()
    messages.success(request, "Associação removida com sucesso.")
    return HttpResponse("")


# =============================================================================
# MATRIZ DE EPI POR FUNÇÃO
# =============================================================================

class ControleEPIPorFuncaoView(SSTBaseMixin, TemplateView):
    template_name = 'seguranca_trabalho/controle_epi_por_funcao.html'
    permission_required = 'seguranca_trabalho.view_matrizepi'

    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        filial_atual = self.get_filial_ativa()

        if not filial_atual:
            messages.warning(self.request, "Nenhuma filial selecionada.")
            context.update({
                'titulo_pagina': "Matriz de EPIs",
                'funcoes': [],
                'equipamentos': [],
                'matriz_data': {},
            })
            return context

        funcoes = Funcao.objects.filter(filial=filial_atual, ativo=True).order_by('nome')
        equipamentos = Equipamento.objects.filter(ativo=True).order_by('nome')
        dados_salvos = (
            MatrizEPI.objects
            .filter(funcao__in=funcoes)
            .select_related('funcao', 'equipamento')
        )

        matriz_data = {}
        for item in dados_salvos:
            matriz_data.setdefault(item.funcao_id, {})
            matriz_data[item.funcao_id][item.equipamento_id] = item.frequencia_troca_meses

        context.update({
            'titulo_pagina': f"Matriz de EPIs - {filial_atual.nome}",
            'funcoes': funcoes,
            'equipamentos': equipamentos,
            'matriz_data': matriz_data,
        })
        return context

    @transaction.atomic
    def post(self, request, *args, **kwargs):
        filial_id = self.get_filial_ativa_id()
        if not filial_id:
            messages.error(request, "Sessão expirada ou filial não selecionada.")
            return redirect(request.path_info)

        funcoes_ids = set(
            Funcao.objects.filter(filial_id=filial_id, ativo=True)
            .values_list('id', flat=True)
        )
        equipamentos_ids = set(
            Equipamento.objects.filter(ativo=True).values_list('id', flat=True)
        )

        existentes = MatrizEPI.objects.filter(funcao_id__in=funcoes_ids)
        mapa_existentes = {
            (item.funcao_id, item.equipamento_id): item
            for item in existentes
        }

        entries_to_create = []
        entries_to_update = []
        submitted_keys = set()

        for key, value in request.POST.items():
            if not key.startswith('freq_'):
                continue
            try:
                _, funcao_id_str, equipamento_id_str = key.split('_')
                funcao_id = int(funcao_id_str)
                equipamento_id = int(equipamento_id_str)
            except (ValueError, IndexError):
                continue

            if funcao_id not in funcoes_ids or equipamento_id not in equipamentos_ids:
                continue

            submitted_keys.add((funcao_id, equipamento_id))
            frequencia = int(value) if value and value.isdigit() else 0
            if frequencia <= 0:
                continue

            chave = (funcao_id, equipamento_id)
            if chave in mapa_existentes:
                item_existente = mapa_existentes[chave]
                if item_existente.frequencia_troca_meses != frequencia:
                    item_existente.frequencia_troca_meses = frequencia
                    entries_to_update.append(item_existente)
            else:
                entries_to_create.append(
                    MatrizEPI(
                        funcao_id=funcao_id,
                        equipamento_id=equipamento_id,
                        filial_id=filial_id,
                        frequencia_troca_meses=frequencia,
                    )
                )

        keys_to_delete = set(mapa_existentes.keys()) - submitted_keys
        pks_to_delete = [mapa_existentes[key].pk for key in keys_to_delete]

        if entries_to_create:
            MatrizEPI.objects.bulk_create(entries_to_create)
        if entries_to_update:
            MatrizEPI.objects.bulk_update(entries_to_update, ['frequencia_troca_meses'])
        if pks_to_delete:
            MatrizEPI.objects.filter(pk__in=pks_to_delete).delete()

        messages.success(request, "Matriz de EPIs salva com sucesso!")
        return redirect(request.path_info)


# =============================================================================
# RELATÓRIOS PDF / WORD - GERAIS
# =============================================================================

class RelatorioSSTPDFView(SSTBaseMixin, View):
    http_method_names = ['get']

    def get(self, request, *args, **kwargs):
        entregas = EntregaEPI.objects.for_request(request).select_related(
            'ficha__funcionario', 'equipamento'
        )
        context = {
            'data_emissao': timezone.now(),
            'entregas': entregas,
            'usuario': request.user,
        }
        html_string = render_to_string(
            'seguranca_trabalho/relatorio_geral_pdf.html', context
        )
        html = HTML(string=html_string, base_url=request.build_absolute_uri())
        pdf = html.write_pdf()
        response = HttpResponse(pdf, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="relatorio_sst.pdf"'
        return response


class ExportarFuncionariosPDFView(SSTBaseMixin, View):
    def get(self, request, *args, **kwargs):
        funcionarios = Funcionario.objects.for_request(request).select_related(
            'cargo', 'departamento'
        ).filter(status='ATIVO')

        context = {
            'funcionarios': funcionarios,
            'data_emissao': timezone.now().strftime('%d/%m/%Y às %H:%M'),
            'nome_filial': self.get_filial_ativa() or 'Geral',
        }
        html_string = render_to_string(
            'departamento_pessoal/relatorio_funcionarios_pdf.html', context
        )
        html = HTML(string=html_string, base_url=request.build_absolute_uri())
        pdf_file = html.write_pdf()

        response = HttpResponse(pdf_file, content_type='application/pdf')
        response['Content-Disposition'] = 'attachment; filename="relatorio_funcionarios.pdf"'
        return response


class ExportarFuncionariosWordView(SSTBaseMixin, View):
    def get(self, request, *args, **kwargs):
        funcionarios = Funcionario.objects.for_request(request).select_related(
            'cargo', 'departamento'
        ).filter(status='ATIVO')

        document = Document()
        document.add_heading('Relatório de Colaboradores', level=1)
        document.add_paragraph(
            f'Gerado em: {timezone.now().strftime("%d/%m/%Y às %H:%M")}'
        )
        document.add_paragraph()

        table = document.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Nome'
        hdr[1].text = 'Cargo'
        hdr[2].text = 'Departamento'
        hdr[3].text = 'Admissão'

        for f in funcionarios:
            row = table.add_row().cells
            row[0].text = f.nome_completo
            row[1].text = f.cargo.nome if f.cargo else '-'
            row[2].text = f.departamento.nome if f.departamento else '-'
            row[3].text = f.data_admissao.strftime('%d/%m/%Y') if f.data_admissao else '-'

        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)

        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = 'attachment; filename="relatorio_funcionarios.docx"'
        return response


# =============================================================================
# REDIRECTS / UTILITÁRIOS
# =============================================================================

@login_required
def minha_ficha_redirect_view(request):
    """Redireciona o usuário para sua própria ficha de EPI."""
    funcionario_obj = Funcionario.objects.filter(usuario=request.user).first()
    if not funcionario_obj:
        messages.error(
            request,
            "Seu usuário não está associado a um perfil de funcionário."
        )
        return redirect('usuario:profile')

    ficha = FichaEPI.objects.filter(funcionario=funcionario_obj).first()
    if not ficha:
        messages.error(request, "Você não possui uma ficha de EPI.")
        return redirect('usuario:profile')

    return redirect('seguranca_trabalho:ficha_detail', pk=ficha.pk)


